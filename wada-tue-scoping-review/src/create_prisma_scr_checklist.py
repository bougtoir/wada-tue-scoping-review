#!/usr/bin/env python3
"""Generate PRISMA-ScR Checklist as a .docx file."""
import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(SCRIPT_DIR, '..', 'manuscripts')
os.makedirs(OUT_DIR, exist_ok=True)

# PRISMA-ScR items per Tricco et al. 2018
# (Section, Item#, Checklist item, Reported on page)
CHECKLIST = [
    ('TITLE', '1', 'Identify the report as a scoping review.', 'Title page'),
    ('ABSTRACT', '', '', ''),
    ('  Structured summary', '2',
     'Provide a structured summary that includes (as applicable): background, objectives, eligibility criteria, sources of evidence, charting methods, results, and conclusions.',
     'Abstract (p. 1)'),
    ('INTRODUCTION', '', '', ''),
    ('  Rationale', '3',
     'Describe the rationale for the review in the context of what is already known. Explain why the review questions/objectives lend themselves to a scoping review approach.',
     'Introduction (pp. 2\u20133)'),
    ('  Objectives', '4',
     'Provide an explicit statement of the questions and objectives being addressed with reference to their key elements (e.g., population or participants, concepts, and context) or other relevant items appropriate for scoping reviews.',
     'Introduction (p. 3)'),
    ('METHODS', '', '', ''),
    ('  Protocol and registration', '5',
     'Indicate whether a review protocol exists; state if and where it can be accessed (e.g., a Web address); and if available, provide registration information, including the registration number.',
     'Methods \u2013 Protocol and Registration'),
    ('  Eligibility criteria', '6',
     'Specify characteristics of the sources of evidence used as eligibility criteria (e.g., years considered, language, and publication status), and provide a rationale.',
     'Methods \u2013 Eligibility Criteria'),
    ('  Information sources', '7',
     'Describe all information sources in the search (e.g., databases with dates of coverage and contact with authors of studies to identify additional sources), as well as the date the most recent search was executed.',
     'Methods \u2013 Information Sources'),
    ('  Search', '8',
     'Present the full search strategy for at least one electronic database, including any limits used, such that it could be repeated.',
     'Methods \u2013 Search Strategy'),
    ('  Selection of sources of evidence', '9',
     'State the process for selecting sources of evidence (i.e., screening and eligibility) included in the scoping review.',
     'Methods \u2013 Selection of Sources'),
    ('  Data charting process', '10',
     'Describe the methods of charting data from the included sources of evidence (e.g., calibrated forms or forms that have been tested by the team before their use, and whether data charting was done independently or in duplicate) and any processes for obtaining and confirming data from investigators.',
     'Methods \u2013 Data Charting'),
    ('  Data items', '11',
     'List and define all variables for which data were sought and any assumptions and simplifications made.',
     'Methods \u2013 Data Charting'),
    ('  Critical appraisal of individual sources of evidence', '12',
     'If done, provide a rationale for conducting a critical appraisal of included sources of evidence; describe the methods used and how this information was used in any data synthesis (if applicable).',
     'Not applicable \u2013 consistent with scoping review methodology'),
    ('  Synthesis of results', '13',
     'Describe the methods of handling and summarizing the data that were charted.',
     'Methods \u2013 Data Charting'),
    ('RESULTS', '', '', ''),
    ('  Selection of sources of evidence', '14',
     'Give numbers of sources of evidence screened, assessed for eligibility, and included in the review, with reasons for exclusions at each stage, ideally using a flow diagram.',
     'Results (Figure 1)'),
    ('  Characteristics of sources of evidence', '15',
     'For each source of evidence, present characteristics for which data were charted and provide the citations.',
     'Results \u2013 Disease-specific sections; Table 1'),
    ('  Critical appraisal within sources of evidence', '16',
     'If done, present data on critical appraisal of included sources of evidence (see item 12).',
     'Not applicable'),
    ('  Results of individual sources of evidence', '17',
     'For each included source of evidence, present the relevant data that were charted that relate to the review questions and objectives.',
     'Results \u2013 Disease-specific sections'),
    ('  Synthesis of results', '18',
     'Summarize and/or present the charting results as they relate to the review questions and objectives.',
     'Results \u2013 Cross-Cutting Analysis; Figures 2\u20134'),
    ('DISCUSSION', '', '', ''),
    ('  Summary of evidence', '19',
     'Summarize the main results (including an overview of concepts, themes, and types of evidence available), link to the review questions and objectives, and consider the relevance to key groups.',
     'Discussion (pp. 10\u201313)'),
    ('  Limitations', '20',
     'Discuss the limitations of the scoping review process.',
     'Discussion \u2013 Limitations'),
    ('  Conclusions', '21',
     'Provide a general interpretation of the results with respect to the review questions and objectives, as well as potential implications and/or next steps.',
     'Conclusions'),
    ('FUNDING', '22',
     'Describe sources of funding for the included sources of evidence, as well as sources of funding for the scoping review. Describe the role of the funders of the scoping review.',
     'Title page \u2013 [No external funding]'),
]


def create_checklist():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'PRISMA-ScR (Preferred Reporting Items for Systematic Reviews and '
        'Meta-Analyses extension for Scoping Reviews) Checklist'
    )
    run.bold = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        'Tricco AC, Lillie E, Zarin W, et al. PRISMA Extension for Scoping Reviews '
        '(PRISMA-ScR): Checklist and Explanation. Ann Intern Med. 2018;169(7):467\u2013473.'
    )
    run.font.size = Pt(9)
    run.italic = True

    doc.add_paragraph()

    # Create table
    table = doc.add_table(rows=1, cols=4, style='Table Grid')
    table.autofit = True

    # Header row
    headers = ['SECTION', 'ITEM', 'PRISMA-ScR CHECKLIST ITEM', 'REPORTED ON PAGE #']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)

    # Data rows
    for section, item, desc, page in CHECKLIST:
        row = table.add_row()
        cells = row.cells
        
        # Section name
        run = cells[0].paragraphs[0].add_run(section.strip())
        run.font.size = Pt(9)
        if not item:  # Section header
            run.bold = True

        # Item number
        run = cells[1].paragraphs[0].add_run(item)
        run.font.size = Pt(9)

        # Description
        run = cells[2].paragraphs[0].add_run(desc)
        run.font.size = Pt(9)

        # Page
        run = cells[3].paragraphs[0].add_run(page)
        run.font.size = Pt(9)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(1.2)
        row.cells[2].width = Cm(10.0)
        row.cells[3].width = Cm(4.0)

    path = os.path.join(OUT_DIR, 'PRISMA_ScR_Checklist.docx')
    doc.save(path)
    print(f'Saved: {path}')


if __name__ == '__main__':
    create_checklist()
