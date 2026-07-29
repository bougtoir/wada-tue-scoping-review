#!/usr/bin/env python3
"""Build point-by-point Response to Reviewers for SCJ revision."""
import os, sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from create_scj_review_part1 import setup_styles, add_heading_styled

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(BASE_DIR, 'manuscripts')
os.makedirs(OUT_DIR, exist_ok=True)

doc = Document()
setup_styles(doc)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Response to Reviewers')
run.bold = True; run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('SCJ-D-26-00070: Mind the Gap: A Scoping Review of Discrepancies Between WADA Therapeutic Use Exemptions and Current Clinical Practice Guidelines')
run.font.size = Pt(12); run.italic = True

doc.add_paragraph()

add_heading_styled(doc, 'General Response', 1)

for para in [
    'We thank the Associate Editor and both reviewers for the constructive and detailed feedback. We have revised the manuscript to address every comment. Major structural changes include: (1) adding page numbers; (2) moving figures and their legends to the end of the manuscript (no embedded figures); (3) moving Source Selection from Results to Methods; (4) placing Recommendations for Harmonization within the Discussion; (5) shortening Practical Applications; and (6) removing the separate Conclusion section. We have also substantially revised the content to provide population/epidemiological context, include WADA perspective throughout, add FDA/EMA/PMDA approval information, explain organic versus functional hypogonadism, address the heart failure/elite-sport issue, and replace PCOS terminology with PMOS while retaining PCOS in parentheses at first use.',
    'All references have been re-verified, renumbered, and alphabetized by lead author per SCJ guidelines. The PRISMA-ScR checklist has been updated and is submitted as a supplementary file. A clean revised manuscript, a marked manuscript (changes shown in red/highlight), and this point-by-point response are provided.',
]:
    pr = doc.add_paragraph()
    run = pr.add_run(para)
    run.font.size = Pt(12)


def add_comment(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('Comment: ')
    run.bold = True; run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.italic = True


def add_response(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('Response: ')
    run.bold = True; run.font.size = Pt(11)
    run = p.add_run(text)
    run.font.size = Pt(11)


add_heading_styled(doc, 'Reviewer #1', 1)

comments_r1 = [
    ('Please include page numbers on the resubmission.',
     'Page numbers have been added to the footer of every page.'),
    ('This scoping review is missing a key element: population (epidemiological data) suggesting that there are people with these diseases/health issues in tested sports.',
     'A new subsection, Population and Epidemiological Context, has been added to Results. It reports 2023 WADA testing figures (288,865 samples across 94 sports/disciplines), athlete asthma prevalence (~20–25%), ADHD prevalence among elite athletes (7–8%), and the expanding masters athlete population (2024 WMA Championships: 8,029 athletes; 2027 World Masters Games: >24,000 registrants).'),
    ('It seems that the role/responsibilities of S&C professionals would be much lower than that of medical staff... the recommendations at the end are very hard to support.',
     'We have reduced the Practical Applications section to four concise, actionable items and reframed them as support/communication roles rather than clinical decision-making. Recommendations for policy harmonization now appear within the Discussion and are directed at WADA, regulatory bodies, and medical stakeholders, not S&C professionals.'),
    ('It is not surprising there are gaps. The updates to clinical guidelines are on different timelines than those of TUE.',
     'This point is now explicitly stated in the Cross-Cutting Analysis and Discussion as the first structural driver: temporal mismatch between clinical guideline update cycles and WADA regulatory revisions.'),
    ('The primary thing missing is FDA/EMA/PMDA guidance... specifically for the newer drugs.',
     'A new paragraph in the Type 2 Diabetes section now lists FDA, EMA, and PMDA approval dates for semaglutide and tirzepatide, with corresponding references.'),
    ('There has been a recent change to PCOS, it is now PMOS - please correct throughout the manuscript.',
     'We have adopted the proposed term PMOS (polyendocrine metabolic ovarian syndrome) throughout, with PCOS retained in parentheses at first use. All figures and legends have also been updated.'),
    ('On page 16, you discuss that WADA is permissive, but WADA cares about performance enhancement, which b2 agonists can do.',
     'The Asthma section now explicitly states WADAs rationale: systemic beta-2 agonists can increase lean body mass and reduce fat mass, and permitted inhaled doses are set below doses associated with these effects. This performance-enhancement rationale is independent of clinical asthma care.'),
    ('On pdf page 16 you mention potential compliance risk. Has this happened in the literature or is there any evidence to support it?',
     'We have qualified the statement as a threshold exceedance risk under acute, clinically appropriate dosing during exacerbations, and note that athletes may apply for retroactive TUEs. We do not claim that documented ADRVs are common; rather, the uncertainty and administrative burden itself may influence treatment decisions.'),
    ('Please provide examples or evidence to support the claim that TUE submissions may be operationally restricted to athletes with elite-level results or prior ADRVs.',
     'We have softened the language and now note that some NADOs require proof of national-team status, international results, or prior anti-doping testing, while others accept applications at any level but impose documentation and processing requirements that function as de facto barriers. We cite athlete-survey and TUE-application literature and explicitly state that systematic eligibility data by competitive tier are limited.'),
    ('T2D section: GLP-1 RAs are not on the list, so this section seems irrelevant.',
     'The section is retained because GLP-1 receptor agonists are on the WADA Monitoring Program and represent an emerging clinical-competition gap; we note that this is not a current TUE issue but a rapidly evolving one.'),
    ('Please provide data on the number of athletes who are >40 in sports monitored by regulatory agencies.',
     'The hypogonadism section now cites a study of 183 male athletes over 50 years showing 11.5% severe and 26.2% mild testosterone deficiency. It also provides masters-sport registration data (WMA 2024; World Masters Games 2027) to contextualize the size of the >40 athlete population.'),
    ('Heart failure? I am fairly certain that is contra-indicated for elite level sport participation.',
     'We now clarify that advanced heart failure generally precludes elite competition, but athletes with earlier-stage or stable heart failure, as well as masters and recreational competitors, may be prescribed MRAs. The example has been revised accordingly.'),
    ('There is no mention of the biological passport.',
     'The Athlete Biological Passport (ABP) and ADAMS are now introduced in the Cross-Cutting Analysis and their role in longitudinal monitoring and TUE administration is discussed.'),
    ('The discussion could be shortened substantially - it is very repetitive with the results.',
     'The Discussion has been shortened and reorganized. It now emphasizes synthesis, WADA perspective, ethical implications, and limitations rather than repeating results.'),
    ('Much like the military and other occupations - it may not be fair that people cannot receive a specific treatment - and that means they cannot compete.',
     'This tension is now acknowledged in the Discussion: anti-doping rules exist because some substances can enhance performance and some athletes misuse them; the clinical-competition gap reflects a necessary trade-off that may limit participation for some athletes.'),
    ('Figure 1 - please identify the number of studies identified in each database.',
     'Figure 1 now lists database-specific counts in the top identification box (PubMed n=312, Embase n=198, SPORTDiscus n=121, Cochrane n=84, Web of Science n=132, Total n=847), and the legend provides these counts.'),
    ('Figure 3 - please explain what the gap sensitivity means more clearly.',
     'The Figure 3 legend has been expanded to state that the timeline shows how clinical guideline updates occur on annual or near-annual cycles while WADA TUE guidance and Prohibited List changes may lag by 1–3 years, creating a window of divergence.'),
    ('Figure 5 - adjust your x-axis so the names can be fully seen.',
     'Figure 5 disease-area labels have been rewritten without line breaks, PCOS has been changed to PMOS, and the x-axis labels are rotated 30 degrees and right-aligned so the full text is visible.'),
    ('Make sure the prisma checklist is included.',
     'An updated PRISMA-ScR checklist is included as a supplementary file and generated from src/create_prisma_scr_checklist.py.'),
]

for c, r in comments_r1:
    add_comment(doc, c)
    add_response(doc, r)
    doc.add_paragraph()

add_heading_styled(doc, 'Reviewer #2', 1)

comments_r2 = [
    ('Add the WADA perspective to create a balanced narrative.',
     'A dedicated paragraph in the Discussion explicitly acknowledges WADAs mandate to protect clean sport, the role of the Prohibited List, ISTUE, ABP, and ADAMS, and the legitimate rationale for threshold-based and organic-only rules. WADA perspective has also been added to the asthma, ADHD, hypogonadism, cardiovascular, and T2D sections.'),
    ('Figure 4 is confusing. Recommend providing an example in the description of how to interpret the figure.',
     'The Figure 4 caption now includes a concrete example (an adult athlete with ADHD on long-acting methylphenidate who would be required to discontinue the medication in-competition without a TUE). The conceptual framework figure itself retains the same structural layout.'),
    ('Fully spell out the following categories on the X axis of Figure 5: Male..., PCOS/..., and Type 2 Diabetes...',
     'Figure 5 labels are now fully spelled out as "Male Hypogonadism," "PMOS / Fertility," and "Type 2 Diabetes / GLP-1 RA" and are rotated for readability.'),
    ('Page 1, Line 20: Confirm with editor if key words can be words that appear in the article title.',
     'Keywords were selected from the title and body of the manuscript per SCJ guidance; we are prepared to adjust if the Editorial Office prefers.'),
    ('Page 2, Line 10: Add "provided such use does not confer an unfair advantage" to the TUE principle sentence.',
     'The Introduction now states that WADAs objective is to permit appropriate medical care provided the use would produce no additional enhancement beyond a return to normal health and does not confer an unfair advantage.'),
    ('Page 5, Lines 14 & 16: Should National Anti-Doping Organizations and International Federations be capitalized?',
     'We have changed these to lower case ("national anti-doping organizations" and "international federations") while retaining the acronyms NADOs and IFs.'),
    ('Page 6, Line 35 / Page 8, Line 23: Are they exclusively international medical organizations or are some on the national level?',
     'The Methods eligibility criteria now explicitly include both recognized international and national medical organizations.'),
    ('Confirm if an acronym can start a sentence (WADA).',
     'We have revised sentences to avoid starting with WADA as an acronym; the expansion "World Anti-Doping Agency" is used when it begins a sentence.'),
    ('Page 10, Line 21: Confirm wording "...could accumulate 48 μg delivered formoterol daily..."',
     'The wording is retained as an illustrative calculation based on a budesonide/formoterol 200/6 μg MART regimen and the 54 μg WADA formoterol threshold.'),
    ('Page 11, Lines 35-43: Include WADA perspective on why ADHD stimulants are allowed out-of-competition but not in-competition.',
     'The ADHD section now explicitly explains WADAs rationale: stimulants may acutely improve alertness and concentration in competition, whereas out-of-competition treatment continuity is permitted because detection windows are shorter and clinical need persists.'),
    ('Page 13, Line 25: This statement requires a reference to support it.',
     'The statement has been supported with references (23, 32, 34): Overbye and Wagner (23) document how elite athletes experience TUE access as a barrier; Vernec et al. (32) report variation in TUE prevalence across Olympic Games; and WADA\'s ISTUE (34) establishes the criteria that NADOs and IFs implement with variable practical requirements.'),
    ('Page 14: Mention the major increase in GLP-1 medications in overweight/obese populations without diabetes.',
     'The T2D section now notes that semaglutide and tirzepatide are also widely used for weight management, which is the basis for WADA surveillance and potential future prohibition in weight-class sports.'),
    ('Page 15, Lines 31 & 33: WADA perspective on testosterone and risk of supraphysiological enhancement.',
     'The hypogonadism section now explicitly states WADAs concern that testosterone restoration can be abused to produce supraphysiological effects, and explains why the organic-only rule seeks to prevent misuse.'),
    ('Page 16, Line 56: PCOS is now PMOS as of May 12, 2026.',
     'We have adopted PMOS (polyendocrine metabolic ovarian syndrome) throughout, with PCOS in parentheses at first use.'),
    ('Page 17, Lines 43 & 45: Explain the difference between organic and functional hypogonadism.',
     'A new paragraph defines organic hypogonadism (identifiable hypothalamic-pituitary-gonadal pathology) and functional hypogonadism (age-, obesity-, opioid-, or overtraining-related suppression without a discrete structural lesion).'),
    ('Page 19, Line 51: Confirm accuracy of reference (9).',
     'Reference (9) is Hostrup et al., a systematic review and meta-analysis of beta-2 agonist performance-enhancing effects, which we now use appropriately in the asthma and discussion sections.'),
    ('Page 19, Line 59: Reference (8) and sentence are misleading.',
     'We have removed the misleading use of Heuberger & Cohen (8) for beta-2 agonists. The asthma section now uses Hostrup et al. (9) for the performance-enhancement rationale and WADA/Allen references for threshold policy.'),
    ('Page 21, Line 31: International Federations capitalization.',
     'Changed to lower case.'),
    ('Page 23, Line 7: Use S&C instead of strength and conditioning.',
     'All instances have been standardized to S&C after first use.'),
    ('Page 25, Line 49: Spell out ADAMS at first use.',
     'The Anti-Doping Administration and Management System (ADAMS) is now spelled out at first use in the Cross-Cutting Analysis.'),
]

for c, r in comments_r2:
    add_comment(doc, c)
    add_response(doc, r)
    doc.add_paragraph()

add_heading_styled(doc, 'Associate Editor', 1)

comments_ae = [
    ('Page 2 – Move the figures to the end of the paper where the figure legends are located.',
     'All figures have been removed from the main text. The manuscript now cites figures parenthetically in the text and presents detailed figure legends on a separate page after References.'),
    ('ABSTRACT Line 1 - Put "and the" between (WADA) and Therapeutic.',
     'The abstract now reads: "The World Anti-Doping Agency (WADA) and the Therapeutic Use Exemption (TUE) system permit athletes..."'),
    ('INTRODUCTION – You need page numbers.',
     'Page numbers have been added to the footer.'),
    ('Page 2, Line 12 – Suggested purpose statement.',
     'The purpose statement has been revised for clarity to reflect mapping discrepancies between WADA TUE regulations and clinical practice guidelines and providing practical guidance for S&C professionals.'),
    ('RESULTS - Source Selection should be in Methods.',
     'Source Selection has been moved to the Methods section under "Selection of Sources of Evidence."'),
    ('DISCUSSION Line 8 – Change to discrepancies between WADA and TUE.',
     'The Discussion opening now reads "...systematic mapping of discrepancies between WADA TUE regulations and current clinical practice guidelines..."'),
    ('PRACTICAL APPLICATIONS - Reduce this section.',
     'Practical Applications has been shortened from seven long bullets to four concise, action-oriented recommendations.'),
    ('RECOMMENDATIONS FOR HARMONIZATION should be in the Discussion.',
     'Recommendations for Harmonization now appear as a subsection within the Discussion.'),
    ('Omit the Conclusions.',
     'The separate Conclusion section has been removed; concluding points are integrated into the final paragraphs of the Discussion.'),
]

for c, r in comments_ae:
    add_comment(doc, c)
    add_response(doc, r)
    doc.add_paragraph()

output_path = os.path.join(OUT_DIR, 'SCJ_Response_to_Reviewers.docx')
doc.save(output_path)
print(f"Saved: {output_path}")
