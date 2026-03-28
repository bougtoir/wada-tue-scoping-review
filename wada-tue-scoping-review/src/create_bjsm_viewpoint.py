#!/usr/bin/env python3
"""Create BJSM Viewpoint/Editorial .docx files (English + Japanese)."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

OUTPUT_DIR = "/home/ubuntu"
FIG_DIR = "/home/ubuntu/figures"


def setup_styles(doc):
    """Set up document styles."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 2.0  # Double spacing

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def add_title_page(doc, title, authors, affiliations, corresponding, word_count, keywords):
    """Add title page."""
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()  # blank line

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(authors)
    run.font.size = Pt(12)

    doc.add_paragraph()

    # Affiliations
    for aff in affiliations:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(aff)
        run.font.size = Pt(10)
        run.italic = True

    doc.add_paragraph()

    # Corresponding author
    p = doc.add_paragraph()
    run = p.add_run('Corresponding Author: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(corresponding)
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Word count
    p = doc.add_paragraph()
    run = p.add_run(f'Word Count: {word_count}')
    run.font.size = Pt(10)

    # Keywords
    p = doc.add_paragraph()
    run = p.add_run('Keywords: ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(keywords)
    run.font.size = Pt(10)

    doc.add_page_break()


def add_heading_styled(doc, text, level=1):
    """Add a styled heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body_text(doc, text):
    """Add body text paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.size = Pt(12)
    return p


def add_body_text_no_indent(doc, text):
    """Add body text without indent."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    return p


def create_bjsm_english():
    """Create BJSM Viewpoint in English."""
    doc = Document()
    setup_styles(doc)

    # === TITLE PAGE ===
    add_title_page(
        doc,
        title='Mind the Gap: Clinical Practice Guidelines and WADA Therapeutic Use Exemptions Are Diverging\u2014A Call for Harmonization',
        authors='[Author Name], MD, CSCS',
        affiliations=[
            '[Institutional Affiliation]',
            '[Department/Division]',
            '[City, Country]'
        ],
        corresponding='[Author Name], MD, CSCS\n[Address]\nEmail: [email]\nORCID: [ORCID ID]',
        word_count='~1,500 (excluding references and figure legend)',
        keywords='therapeutic use exemption; anti-doping; clinical practice guidelines; prohibited list; athlete health'
    )

    # === MAIN TEXT ===
    add_heading_styled(doc, 'INTRODUCTION', 1)

    add_body_text(doc,
        'The World Anti-Doping Agency (WADA) Therapeutic Use Exemption (TUE) system represents '
        'a critical safeguard, ensuring that athletes with legitimate medical conditions can receive '
        'necessary treatment while preserving sporting integrity. However, an emerging and under-recognized '
        'problem threatens this balance: the growing divergence between WADA\u2019s regulatory framework and '
        'current evidence-based clinical practice guidelines. We argue that this \u201cclinical-competition gap\u201d '
        'places athletes at risk of receiving suboptimal medical care or inadvertent anti-doping rule violations (ADRVs).'
    )

    add_heading_styled(doc, 'THE WIDENING GAP', 1)

    add_body_text(doc,
        'We conducted a scoping review mapping WADA\u2019s 2025\u20132026 Prohibited List and TUE Physician '
        'Guidelines against current clinical practice guidelines across seven major disease areas. The findings '
        'reveal systematic discrepancies of clinical significance (Figure 1).'
    )

    add_body_text(doc,
        'In male hypogonadism, the Endocrine Society (2018) and EAU (2024) guidelines recommend testosterone '
        'replacement therapy (TRT) for men with consistently low testosterone and symptoms, regardless of whether '
        'the etiology is \u201corganic\u201d or \u201cfunctional.\u201d The TRAVERSE trial (2023) further established cardiovascular '
        'safety of TRT. Yet WADA\u2019s TUE Physician Guidelines explicitly restrict TUE approval to organic causes, '
        'excluding late-onset hypogonadism (LOH)\u2014precisely the condition affecting the rapidly growing masters '
        'athlete population. An athlete with biochemically confirmed LOH (testosterone <300 ng/dL on two morning '
        'measurements) whose symptoms persist despite lifestyle modification faces a binary choice: forgo '
        'guideline-recommended treatment or retire from competition.'
    )

    add_body_text(doc,
        'For attention-deficit/hyperactivity disorder (ADHD), all major guidelines\u2014including NICE (2024), '
        'the Australian ADHD Guideline (2022), and APA recommendations\u2014designate stimulant medications '
        '(methylphenidate, lisdexamfetamine) as first-line pharmacotherapy. These are universally prohibited '
        'in-competition under WADA\u2019s S6 category. The paradox is stark: ADHD is a 24/7 condition, yet WADA '
        'prohibits its optimal treatment only in-competition, effectively requiring athletes to discontinue '
        'long-acting formulations on competition days\u2014a practice no clinical guideline endorses. While TUE '
        'applications for stimulants are possible, the process demands extensive neuropsychological documentation '
        'dating to childhood, imposing substantial time and financial burdens.'
    )

    add_body_text(doc,
        'The asthma landscape presents a subtler but significant challenge. GINA 2025 established '
        'ICS-formoterol maintenance and reliever therapy (MART) as the preferred track across all severity '
        'steps. During exacerbations, cumulative daily formoterol doses may approach WADA\u2019s 54 \u03bcg '
        'threshold (delivered dose per 24 hours), and the urinary formoterol threshold of 40 ng/mL may be '
        'exceeded with clinically appropriate MART use\u2014creating ADRV risk for athletes following '
        'guideline-recommended therapy.'
    )

    add_body_text(doc,
        'Perhaps the most pressing emerging gap involves GLP-1 receptor agonists. The ADA Standards of '
        'Care 2025 now positions semaglutide and tirzepatide as preferred agents for type 2 diabetes with '
        'cardiovascular risk or weight management needs. These agents are currently on WADA\u2019s Monitoring '
        'Program (since 2024) and may be added to the Prohibited List. Should this occur, athletes with '
        'type 2 diabetes following ADA guidelines would be forced to abandon first-line therapy, potentially '
        'shifting to insulin\u2014which itself requires a TUE and carries hypoglycemic risk.'
    )

    add_body_text(doc,
        'Additional gaps exist in cardiovascular medicine (mineralocorticoid receptor antagonists such as '
        'spironolactone, a pillar of heart failure therapy, are prohibited as diuretics), glucocorticoid '
        'management (2022 prohibition of intra-articular injections in-competition with washout periods '
        'inconsistent with clinical practice), and PCOS/fertility treatment (letrozole, the WHO-recommended '
        'first-line ovulation induction agent, is always prohibited).'
    )

    add_heading_styled(doc, 'WHY DOES THIS GAP EXIST?', 1)

    add_body_text(doc,
        'Three structural factors drive this divergence. First, a temporal mismatch: clinical guidelines '
        'undergo rapid evidence-based revisions (GINA annually, ADA annually, ESC every 3\u20135 years), while '
        'WADA TUE Physician Guidelines, though described as \u201cliving documents,\u201d may lag by 1\u20133 years in '
        'reflecting new clinical standards. Second, a criteria mismatch: clinical guidelines are grounded in '
        'patient-centered outcomes, whereas WADA\u2019s criteria emphasize preventing performance enhancement\u2014'
        'sometimes leading to divergent interpretations of the same evidence (as seen in the organic vs. '
        'functional hypogonadism distinction). Third, an emerging drug mismatch: novel therapeutic classes '
        '(GLP-1 receptor agonists, SGLT2 inhibitors) are adopted into clinical practice faster than WADA\u2019s '
        'regulatory apparatus can evaluate and accommodate them.'
    )

    add_heading_styled(doc, 'WHAT SHOULD WE DO?', 1)

    add_body_text(doc,
        'We propose four actionable steps. First, WADA should establish a formal mechanism to review '
        'and update TUE Physician Guidelines within 12 months of major clinical guideline revisions. '
        'Second, the rigid organic/functional distinction for hypogonadism TUEs should be re-evaluated '
        'in light of the Endocrine Society\u2019s symptom-plus-biochemistry approach and the growing masters '
        'athlete population. Third, before adding GLP-1 receptor agonists to the Prohibited List, WADA '
        'should conduct a formal impact assessment on athletes with type 2 diabetes. Fourth, anti-doping '
        'education for prescribing physicians\u2014who often have limited knowledge of the Prohibited List\u2014'
        'should be integrated into sports medicine curricula and continuing medical education.'
    )

    add_body_text(doc,
        'The TUE system was designed to protect athletes\u2019 right to health. When clinical guidelines and '
        'anti-doping regulations diverge, it is athletes who bear the consequences\u2014forced to choose between '
        'optimal medical care and their sporting careers. Closing this gap is not merely a regulatory '
        'convenience; it is an ethical imperative.'
    )

    doc.add_page_break()

    # === FIGURE ===
    add_heading_styled(doc, 'FIGURE LEGEND', 1)
    p = doc.add_paragraph()
    run = p.add_run('Figure 1. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        'Conceptual framework illustrating the clinical-competition gap in anti-doping. '
        'Evidence-based clinical practice guidelines and WADA anti-doping regulations converge on '
        'the athlete, creating gaps through update timing lags, divergent diagnostic criteria, '
        'prohibition of guideline-recommended first-line therapies, and TUE process barriers. '
        'These gaps result in suboptimal treatment, delayed care access, anti-doping rule violation '
        'risk, privacy concerns, geographic disparities, and career impact.'
    )
    run.font.size = Pt(10)

    doc.add_paragraph()
    fig_path = os.path.join(FIG_DIR, 'fig4_conceptual_framework.png')
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # === REFERENCES ===
    add_heading_styled(doc, 'REFERENCES', 1)

    references = [
        '1. World Anti-Doping Agency. World Anti-Doping Code International Standard for Therapeutic Use Exemptions (ISTUE). Montreal: WADA; 2021.',
        '2. World Anti-Doping Agency. The 2026 Prohibited List International Standard. Montreal: WADA; 2025.',
        '3. World Anti-Doping Agency. TUE Physician Guidelines \u2013 Male Hypogonadism. Montreal: WADA; 2025.',
        '4. Bhasin S, Brito JP, Cunningham GR, et al. Testosterone therapy in men with hypogonadism: an Endocrine Society clinical practice guideline. J Clin Endocrinol Metab. 2018;103(5):1715\u20131744.',
        '5. Lincoff AM, Bhasin S, Flevaris P, et al. Cardiovascular safety of testosterone-replacement therapy. N Engl J Med. 2023;389(2):107\u2013117.',
        '6. European Association of Urology. EAU Guidelines on Sexual and Reproductive Health. Arnhem: EAU; 2024.',
        '7. National Institute for Health and Care Excellence. Attention deficit hyperactivity disorder: diagnosis and management. NICE guideline [NG87]. London: NICE; 2024 (updated).',
        '8. Australian ADHD Professionals Association. Australian Evidence-Based Clinical Practice Guideline for ADHD. Melbourne: AADPA; 2022.',
        '9. World Anti-Doping Agency. TUE Physician Guidelines \u2013 ADHD. Version 8.0. Montreal: WADA; 2026.',
        '10. Global Initiative for Asthma. Global Strategy for Asthma Management and Prevention. Fontana: GINA; 2025.',
        '11. World Anti-Doping Agency. TUE Physician Guidelines \u2013 Asthma. Version 9.3. Montreal: WADA; 2026.',
        '12. American Diabetes Association Professional Practice Committee. Standards of Care in Diabetes\u20142025. Diabetes Care. 2025;48(Suppl 1):S1\u2013S352.',
        '13. World Anti-Doping Agency. 2026 Monitoring Program. Montreal: WADA; 2025.',
        '14. McDonagh TA, Metra M, Adamo M, et al. 2021 ESC Guidelines for the diagnosis and treatment of acute and chronic heart failure. Eur Heart J. 2021;42(36):3599\u20133726.',
        '15. Teede HJ, Tay CT, Laven JJE, et al. Recommendations from the 2023 international evidence-based guideline for the assessment and management of polycystic ovary syndrome. J Clin Endocrinol Metab. 2023;108(10):2447\u20132469.',
        '16. Vernec A, Healy D, Banon T, Petroczi A. Prevalence of therapeutic use exemptions at the Olympic Games and Paralympic Games: an analysis of data from 2016 to 2022. Br J Sports Med. 2024;58(17):966\u2013972.',
        '17. World Anti-Doping Agency. Glucocorticoids and Therapeutic Use Exemptions Guidelines. Montreal: WADA; 2025.',
        '18. Heuberger JAAC, Cohen AF. Review of WADA prohibited substances: limited evidence for performance-enhancing effects. Sports Med. 2019;49(4):525\u2013539.',
    ]

    for ref in references:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 2.0

    # Save
    output_path = os.path.join(OUTPUT_DIR, 'BJSM_Viewpoint_English.docx')
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


def create_bjsm_japanese():
    """Create BJSM Viewpoint in Japanese."""
    doc = Document()
    setup_styles(doc)

    # Title page
    add_title_page(
        doc,
        title='Mind the Gap: \u81e8\u5e8a\u8a3a\u7642\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u3068WADA\u6cbb\u7642\u4f7f\u7528\u7279\u4f8b\u306e\u4e56\u96e2\u2014\u8abf\u548c\u3078\u306e\u63d0\u8a00',
        authors='[\u8457\u8005\u540d], MD, CSCS',
        affiliations=[
            '[\u6240\u5c5e\u6a5f\u95a2]',
            '[\u90e8\u7f72/\u8a3a\u7642\u79d1]',
            '[\u90fd\u5e02, \u56fd]'
        ],
        corresponding='[\u8457\u8005\u540d], MD, CSCS\n[\u4f4f\u6240]\nEmail: [email]\nORCID: [ORCID ID]',
        word_count='\u7d041,500\u8a9e\uff08\u82f1\u8a9e\u7248\u63db\u7b97\u3001\u53c2\u8003\u6587\u732e\u30fb\u56f3\u8868\u9664\u304f\uff09',
        keywords='therapeutic use exemption; anti-doping; clinical practice guidelines; prohibited list; athlete health'
    )

    # Main text
    add_heading_styled(doc, '\u306f\u3058\u3081\u306b', 1)

    add_body_text(doc,
        '\u4e16\u754c\u30a2\u30f3\u30c1\u30fb\u30c9\u30fc\u30d4\u30f3\u30b0\u6a5f\u69cb\uff08WADA\uff09\u306e\u6cbb\u7642\u4f7f\u7528\u7279\u4f8b\uff08TUE: Therapeutic Use Exemption\uff09'
        '\u5236\u5ea6\u306f\u3001\u6b63\u5f53\u306a\u533b\u7642\u30cb\u30fc\u30ba\u3092\u6301\u3064\u30a2\u30b9\u30ea\u30fc\u30c8\u304c\u5fc5\u8981\u306a\u6cbb\u7642\u3092\u53d7\u3051\u306a\u304c\u3089\u7af6\u6280\u306e\u516c\u6b63\u6027\u3092\u4fdd\u3064\u305f\u3081\u306e\u91cd\u8981\u306a\u5236\u5ea6\u3067\u3042\u308b\u3002'
        '\u3057\u304b\u3057\u306a\u304c\u3089\u3001\u3053\u306e\u5236\u5ea6\u306e\u30d0\u30e9\u30f3\u30b9\u3092\u8105\u304b\u3059\u554f\u984c\u304c\u9855\u5728\u5316\u3057\u3064\u3064\u3042\u308b\u3002\u305d\u308c\u306f\u3001WADA\u306e\u898f\u5236\u67a0\u7d44\u307f\u3068\u6700\u65b0\u306e\u30a8\u30d3\u30c7\u30f3\u30b9\u306b\u57fa\u3065\u304f\u81e8\u5e8a\u8a3a\u7642\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3'
        '\u3068\u306e\u9593\u306b\u751f\u3058\u308b\u300c\u81e8\u5e8a\u30fb\u7af6\u6280\u30ae\u30e3\u30c3\u30d7\uff08clinical-competition gap\uff09\u300d\u3067\u3042\u308b\u3002\u3053\u306e\u4e56\u96e2\u306b\u3088\u308a\u3001\u30a2\u30b9\u30ea\u30fc\u30c8\u306f'
        '\u6b21\u5584\u306e\u533b\u7642\u3092\u53d7\u3051\u3089\u308c\u306a\u3044\u30ea\u30b9\u30af\u3084\u3001\u610f\u56f3\u305b\u305a\u30a2\u30f3\u30c1\u30fb\u30c9\u30fc\u30d4\u30f3\u30b0\u898f\u5247\u9055\u53cd\uff08ADRV\uff09\u3092\u72af\u3059\u30ea\u30b9\u30af\u306b\u3055\u3089\u3055\u308c\u308b\u3002'
    )

    add_heading_styled(doc, '\u62e1\u5927\u3059\u308b\u30ae\u30e3\u30c3\u30d7', 1)

    add_body_text(doc,
        '\u6211\u3005\u306f\u3001WADA\u306e2025\u20132026\u5e74\u7981\u6b62\u8868\u304a\u3088\u3073TUE\u533b\u5e2b\u5411\u3051\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u3092\u30017\u3064\u306e\u4e3b\u8981\u75be\u60a3\u9818\u57df\u306e\u6700\u65b0\u81e8\u5e8a\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u3068\u7167\u5408\u3059\u308b'
        '\u30b9\u30b3\u30fc\u30d4\u30f3\u30b0\u30ec\u30d3\u30e5\u30fc\u3092\u5b9f\u65bd\u3057\u305f\u3002\u305d\u306e\u7d50\u679c\u3001\u81e8\u5e8a\u7684\u306b\u91cd\u8981\u306a\u4f53\u7cfb\u7684\u4e56\u96e2\u304c\u660e\u3089\u304b\u3068\u306a\u3063\u305f\uff08\u56f31\uff09\u3002'
    )

    add_body_text(doc,
        '\u7537\u6027\u6027\u816b\u6a5f\u80fd\u4f4e\u4e0b\u75c7\u3067\u306f\u3001Endocrine Society\uff082018\uff09\u304a\u3088\u3073EAU\uff082024\uff09\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u304c\u3001\u75c7\u72b6\u3068\u751f\u5316\u5b66\u7684\u78ba\u8a8d\u306b\u57fa\u3065\u304d\u3001'
        '\u75c5\u56e0\u304c\u300c\u5668\u8cea\u7684\u300d\u304b\u300c\u6a5f\u80fd\u7684\u300d\u304b\u3092\u554f\u308f\u305a\u30c6\u30b9\u30c8\u30b9\u30c6\u30ed\u30f3\u88dc\u5145\u7642\u6cd5\uff08TRT\uff09\u3092\u63a8\u5968\u3057\u3066\u3044\u308b\u3002TRAVERSE\u8a66\u9a13\uff082023\uff09\u306fTRT\u306e'
        '\u5fc3\u8840\u7ba1\u5b89\u5168\u6027\u3092\u5b9f\u8a3c\u3057\u305f\u3002\u3057\u304b\u3057WADA\u306eTUE\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u306f\u5668\u8cea\u7684\u75c5\u56e0\u306b\u306e\u307fTUE\u627f\u8a8d\u3092\u9650\u5b9a\u3057\u3001\u9045\u767a\u6027\u6027\u816b\u6a5f\u80fd\u4f4e\u4e0b\u75c7\uff08LOH\uff09\u3092'
        '\u9664\u5916\u3057\u3066\u3044\u308b\u3002\u3053\u308c\u306f\u307e\u3055\u306b\u3001\u6025\u589e\u3059\u308b\u30de\u30b9\u30bf\u30fc\u30ba\u30a2\u30b9\u30ea\u30fc\u30c8\u4eba\u53e3\u304c\u6700\u3082\u5f71\u97ff\u3092\u53d7\u3051\u308b\u75be\u60a3\u3067\u3042\u308b\u3002'
    )

    add_body_text(doc,
        'ADHD\u3067\u306f\u3001NICE\uff082024\uff09\u3001\u30aa\u30fc\u30b9\u30c8\u30e9\u30ea\u30a2ADHD\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\uff082022\uff09\u7b49\u306e\u5168\u3066\u306e\u4e3b\u8981\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u304c\u523a\u6fc0\u85ac\uff08\u30e1\u30c1\u30eb\u30d5\u30a7\u30cb\u30c7\u30fc\u30c8\u3001'
        '\u30ea\u30b9\u30c7\u30ad\u30b5\u30f3\u30d5\u30a7\u30bf\u30df\u30f3\uff09\u3092\u7b2c\u4e00\u9078\u629e\u85ac\u3068\u3059\u308b\u304c\u3001\u3053\u308c\u3089\u306fWADA S6\u30ab\u30c6\u30b4\u30ea\u30fc\u3067\u7af6\u6280\u4f1a\u6642\u7981\u6b62\u3067\u3042\u308b\u3002'
        'ADHD\u306f24\u6642\u9593365\u65e5\u306e\u969c\u5bb3\u3067\u3042\u308a\u3001\u7af6\u6280\u65e5\u306e\u307f\u670d\u85ac\u3092\u4e2d\u65ad\u3059\u308b\u3053\u3068\u3092\u63a8\u5968\u3059\u308b\u81e8\u5e8a\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u306f\u5b58\u5728\u3057\u306a\u3044\u3002'
    )

    add_body_text(doc,
        '\u5598\u606f\u3067\u306fGINA 2025\u304cICS-\u30db\u30eb\u30e2\u30c6\u30ed\u30fc\u30eb\u306eMART\u7642\u6cd5\u3092\u63a8\u5968\u30c8\u30e9\u30c3\u30af\u3068\u3057\u3001\u589e\u60aa\u6642\u306e\u7d2f\u7a4d\u7528\u91cf\u304cWADA\u95be\u503c\uff0854\u03bcg/24\u6642\u9593\uff09\u306b'
        '\u8fd1\u63a5\u3059\u308b\u30ea\u30b9\u30af\u304c\u3042\u308b\u3002GLP-1\u53d7\u5bb9\u4f53\u4f5c\u52d5\u85ac\u306fADA 2025\u304c2\u578b\u7cd6\u5c3f\u75c5\u306e\u6a19\u6e96\u6cbb\u7642\u3068\u3057\u3066\u5f37\u304f\u63a8\u5968\u3059\u308b\u304c\u3001'
        'WADA\u30e2\u30cb\u30bf\u30ea\u30f3\u30b0\u30d7\u30ed\u30b0\u30e9\u30e0\u306b\u53ce\u8f09\u4e2d\u3067\u3042\u308a\u3001\u5c06\u6765\u7684\u7981\u6b62\u306e\u53ef\u80fd\u6027\u304c\u3042\u308b\u3002\u307e\u305f\u3001\u5fc3\u4e0d\u5168\u6cbb\u7642\u306e4\u672c\u67f1\u306e\u4e00\u3064\u3067\u3042\u308b'
        'MRA\uff08\u30b9\u30d4\u30ed\u30ce\u30e9\u30af\u30c8\u30f3\uff09\u304c\u5229\u5c3f\u85ac\u3068\u3057\u3066\u7981\u6b62\u3001PCOS/\u4e0d\u5984\u6cbb\u7642\u306e\u7b2c\u4e00\u9078\u629e\u85ac\u30ec\u30c8\u30ed\u30be\u30fc\u30eb\u304c\u5e38\u6642\u7981\u6b62\u306a\u3069\u3001'
        '\u8907\u6570\u9818\u57df\u3067\u4e56\u96e2\u304c\u78ba\u8a8d\u3055\u308c\u305f\u3002'
    )

    add_heading_styled(doc, '\u306a\u305c\u30ae\u30e3\u30c3\u30d7\u306f\u751f\u3058\u308b\u306e\u304b', 1)

    add_body_text(doc,
        '\u3053\u306e\u4e56\u96e2\u306e\u80cc\u666f\u306b\u306f3\u3064\u306e\u69cb\u9020\u7684\u8981\u56e0\u304c\u3042\u308b\u3002\u7b2c\u4e00\u306b\u3001\u6642\u9593\u7684\u4e0d\u6574\u5408\uff1a\u81e8\u5e8a\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u306f\u8fc5\u901f\u306b\u6539\u8a02\u3055\u308c\u308b\u304c\u3001'
        'WADA TUE\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u306e\u53cd\u6620\u306b\u306f1\u30123\u5e74\u306e\u30e9\u30b0\u304c\u3042\u308b\u3002\u7b2c\u4e8c\u306b\u3001\u57fa\u6e96\u306e\u4e0d\u6574\u5408\uff1a\u81e8\u5e8a\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u306f\u60a3\u8005\u4e2d\u5fc3\u306e'
        '\u30a2\u30a6\u30c8\u30ab\u30e0\u306b\u57fa\u3065\u304f\u304c\u3001WADA\u306f\u30d1\u30d5\u30a9\u30fc\u30de\u30f3\u30b9\u5411\u4e0a\u306e\u9632\u6b62\u3092\u91cd\u8996\u3059\u308b\u3002\u7b2c\u4e09\u306b\u3001\u65b0\u85ac\u306e\u4e0d\u6574\u5408\uff1a\u65b0\u898f\u6cbb\u7642\u85ac\u304c'
        'WADA\u306e\u898f\u5236\u8a55\u4fa1\u3088\u308a\u3082\u901f\u304f\u81e8\u5e8a\u306b\u5c0e\u5165\u3055\u308c\u308b\u3002'
    )

    add_heading_styled(doc, '\u63d0\u8a00', 1)

    add_body_text(doc,
        '\u6211\u3005\u306f4\u3064\u306e\u5177\u4f53\u7684\u63d0\u8a00\u3092\u884c\u3046\u3002\u7b2c\u4e00\u306b\u3001\u4e3b\u8981\u81e8\u5e8a\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u6539\u8a02\u304b\u308912\u30f6\u6708\u4ee5\u5185\u306eTUE\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u66f4\u65b0\u30e1\u30ab\u30cb\u30ba\u30e0\u306e\u78ba\u7acb\u3002'
        '\u7b2c\u4e8c\u306b\u3001\u6027\u816b\u6a5f\u80fd\u4f4e\u4e0b\u75c7TUE\u306b\u304a\u3051\u308b\u5668\u8cea\u7684/\u6a5f\u80fd\u7684\u533a\u5225\u306e\u518d\u8a55\u4fa1\u3002\u7b2c\u4e09\u306b\u3001GLP-1\u53d7\u5bb9\u4f53\u4f5c\u52d5\u85ac\u306e\u7981\u6b62\u8868\u8ffd\u52a0\u524d\u306e'
        '\u7cd6\u5c3f\u75c5\u30a2\u30b9\u30ea\u30fc\u30c8\u3078\u306e\u5f71\u97ff\u8a55\u4fa1\u5b9f\u65bd\u3002\u7b2c\u56db\u306b\u3001\u51e6\u65b9\u533b\u5411\u3051\u30a2\u30f3\u30c1\u30fb\u30c9\u30fc\u30d4\u30f3\u30b0\u6559\u80b2\u306e\u5f37\u5316\u3002'
    )

    add_body_text(doc,
        'TUE\u5236\u5ea6\u306f\u30a2\u30b9\u30ea\u30fc\u30c8\u306e\u5065\u5eb7\u6a29\u3092\u5b88\u308b\u305f\u3081\u306b\u8a2d\u8a08\u3055\u308c\u305f\u3002\u81e8\u5e8a\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u3068\u30a2\u30f3\u30c1\u30fb\u30c9\u30fc\u30d4\u30f3\u30b0\u898f\u5236\u304c'
        '\u4e56\u96e2\u3059\u308b\u3068\u304d\u3001\u305d\u306e\u4ee3\u511f\u3092\u6255\u3046\u306e\u306f\u30a2\u30b9\u30ea\u30fc\u30c8\u3067\u3042\u308b\u3002\u3053\u306e\u30ae\u30e3\u30c3\u30d7\u3092\u57cb\u3081\u308b\u3053\u3068\u306f\u898f\u5236\u4e0a\u306e\u4fbf\u5b9c\u3067\u306f\u306a\u304f\u3001'
        '\u502b\u7406\u7684\u8cac\u52d9\u3067\u3042\u308b\u3002'
    )

    doc.add_page_break()

    # Figure
    add_heading_styled(doc, '\u56f3\u8868\u8aac\u660e', 1)
    p = doc.add_paragraph()
    run = p.add_run('\u56f31. ')
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(
        '\u30a2\u30f3\u30c1\u30fb\u30c9\u30fc\u30d4\u30f3\u30b0\u306b\u304a\u3051\u308b\u81e8\u5e8a\u30fb\u7af6\u6280\u30ae\u30e3\u30c3\u30d7\u306e\u6982\u5ff5\u7684\u67a0\u7d44\u307f\u3002'
        '\u30a8\u30d3\u30c7\u30f3\u30b9\u306b\u57fa\u3065\u304f\u81e8\u5e8a\u8a3a\u7642\u30ac\u30a4\u30c9\u30e9\u30a4\u30f3\u3068WADA\u30a2\u30f3\u30c1\u30fb\u30c9\u30fc\u30d4\u30f3\u30b0\u898f\u5236\u306e\u4e21\u8005\u304c\u30a2\u30b9\u30ea\u30fc\u30c8\u306b\u53ce\u6582\u3057\u3001'
        '\u66f4\u65b0\u6642\u671f\u306e\u30e9\u30b0\u3001\u8a3a\u65ad\u57fa\u6e96\u306e\u76f8\u9055\u3001\u63a8\u5968\u7b2c\u4e00\u9078\u629e\u85ac\u306e\u7981\u6b62\u3001TUE\u624b\u7d9a\u306e\u969c\u58c1\u3092\u901a\u3058\u3066\u30ae\u30e3\u30c3\u30d7\u3092\u751f\u3058\u3055\u305b\u308b\u3002'
    )
    run.font.size = Pt(10)

    doc.add_paragraph()
    fig_path = os.path.join(FIG_DIR, 'fig4_conceptual_framework.png')
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # References (same as English)
    add_heading_styled(doc, 'REFERENCES', 1)
    references = [
        '1. World Anti-Doping Agency. World Anti-Doping Code International Standard for Therapeutic Use Exemptions (ISTUE). Montreal: WADA; 2021.',
        '2. World Anti-Doping Agency. The 2026 Prohibited List International Standard. Montreal: WADA; 2025.',
        '3. World Anti-Doping Agency. TUE Physician Guidelines \u2013 Male Hypogonadism. Montreal: WADA; 2025.',
        '4. Bhasin S, Brito JP, Cunningham GR, et al. Testosterone therapy in men with hypogonadism: an Endocrine Society clinical practice guideline. J Clin Endocrinol Metab. 2018;103(5):1715\u20131744.',
        '5. Lincoff AM, Bhasin S, Flevaris P, et al. Cardiovascular safety of testosterone-replacement therapy. N Engl J Med. 2023;389(2):107\u2013117.',
        '6. European Association of Urology. EAU Guidelines on Sexual and Reproductive Health. Arnhem: EAU; 2024.',
        '7. National Institute for Health and Care Excellence. Attention deficit hyperactivity disorder: diagnosis and management. NICE guideline [NG87]. London: NICE; 2024.',
        '8. Australian ADHD Professionals Association. Australian Evidence-Based Clinical Practice Guideline for ADHD. Melbourne: AADPA; 2022.',
        '9. World Anti-Doping Agency. TUE Physician Guidelines \u2013 ADHD. Version 8.0. Montreal: WADA; 2026.',
        '10. Global Initiative for Asthma. Global Strategy for Asthma Management and Prevention. Fontana: GINA; 2025.',
        '11. World Anti-Doping Agency. TUE Physician Guidelines \u2013 Asthma. Version 9.3. Montreal: WADA; 2026.',
        '12. American Diabetes Association Professional Practice Committee. Standards of Care in Diabetes\u20142025. Diabetes Care. 2025;48(Suppl 1):S1\u2013S352.',
        '13. World Anti-Doping Agency. 2026 Monitoring Program. Montreal: WADA; 2025.',
        '14. McDonagh TA, Metra M, Adamo M, et al. 2021 ESC Guidelines for the diagnosis and treatment of acute and chronic heart failure. Eur Heart J. 2021;42(36):3599\u20133726.',
        '15. Teede HJ, Tay CT, Laven JJE, et al. Recommendations from the 2023 international evidence-based guideline for the assessment and management of polycystic ovary syndrome. J Clin Endocrinol Metab. 2023;108(10):2447\u20132469.',
        '16. Vernec A, Healy D, Banon T, Petroczi A. Prevalence of therapeutic use exemptions at the Olympic Games and Paralympic Games: an analysis of data from 2016 to 2022. Br J Sports Med. 2024;58(17):966\u2013972.',
        '17. World Anti-Doping Agency. Glucocorticoids and Therapeutic Use Exemptions Guidelines. Montreal: WADA; 2025.',
        '18. Heuberger JAAC, Cohen AF. Review of WADA prohibited substances: limited evidence for performance-enhancing effects. Sports Med. 2019;49(4):525\u2013539.',
    ]
    for ref in references:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        run.font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 2.0

    output_path = os.path.join(OUTPUT_DIR, 'BJSM_Viewpoint_Japanese.docx')
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


if __name__ == '__main__':
    create_bjsm_english()
    create_bjsm_japanese()
    print("BJSM Viewpoint files complete.")
