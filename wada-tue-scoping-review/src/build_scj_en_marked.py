#!/usr/bin/env python3
"""Create a marked revision manuscript by highlighting new/modified paragraphs in red."""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OLD_PATH = '/tmp/old_scj_en.docx'
NEW_PATH = os.path.join(BASE_DIR, 'manuscripts', 'SCJ_Narrative_Review_English.docx')
OUT_PATH = os.path.join(BASE_DIR, 'manuscripts', 'SCJ_Narrative_Review_English_marked.docx')

RED = RGBColor(0xFF, 0x00, 0x00)

old_doc = Document(OLD_PATH)
new_doc = Document(NEW_PATH)

old_texts = set(p.text.strip() for p in old_doc.paragraphs if p.text.strip())

# Insert a cover note at the very beginning
note = new_doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run('MARKED REVISION FILE')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RED
note2 = new_doc.add_paragraph()
note2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = note2.add_run('New and substantially modified text is shown in red. Minor wording, citation renumbering, and reference list changes are not individually marked.')
r2.italic = True; r2.font.size = Pt(11); r2.font.color.rgb = RED
new_doc.add_page_break()

# Move the inserted note/page break to the start of the body element
body = new_doc.element.body
# The last three elements are note, note2, page break; move them to the top
for elem in [new_doc.paragraphs[-3]._element, new_doc.paragraphs[-2]._element, new_doc.paragraphs[-1]._element]:
    body.insert(0, elem)

for para in new_doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    if text not in old_texts:
        for run in para.runs:
            run.font.color.rgb = RED

new_doc.save(OUT_PATH)
print(f"Saved marked version: {OUT_PATH}")
