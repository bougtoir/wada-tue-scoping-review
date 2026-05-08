#!/usr/bin/env python3
"""Generate SCJ Cover Letter (EN + JP) as .docx files."""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(SCRIPT_DIR, '..', 'manuscripts')
os.makedirs(OUT_DIR, exist_ok=True)


def add_para(doc, text, bold=False, italic=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def create_english():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15

    # Date and address
    add_para(doc, '[Date]')
    add_para(doc, '')
    add_para(doc, 'Editor-in-Chief\nStrength and Conditioning Journal\nNational Strength and Conditioning Association')
    add_para(doc, '')
    add_para(doc, 'Dear Editor,', bold=True)

    # Opening
    add_para(doc, (
        'We respectfully submit the enclosed manuscript entitled "Mind the Gap: A Scoping Review of '
        'Discrepancies Between WADA Therapeutic Use Exemption Regulations and Current Clinical Practice '
        'Guidelines" for consideration as a Narrative Review article in the Strength and Conditioning Journal.'
    ))

    # What the paper does
    add_para(doc, (
        'This PRISMA-ScR\u2013compliant scoping review systematically maps discrepancies between the World '
        'Anti-Doping Agency (WADA) Therapeutic Use Exemption (TUE) regulatory framework and current '
        'evidence-based clinical practice guidelines across seven disease areas: asthma, ADHD, type 2 '
        'diabetes, male hypogonadism, glucocorticoid-requiring conditions, cardiovascular disease, and '
        'polycystic ovary syndrome (PCOS). We identified 68 relevant sources and found clinically '
        'significant gaps in all seven areas, with the most critical discrepancies in male hypogonadism, '
        'ADHD, and emerging GLP-1 receptor agonist therapies.'
    ))

    # Why SCJ specifically
    add_para(doc, (
        'We believe this manuscript is particularly well-suited for the SCJ readership for several reasons. '
        'First, the clinical-competition gap directly affects athletes whom strength and conditioning (S&C) '
        'professionals supervise daily. S&C practitioners are often the first to observe performance changes '
        'attributable to medication adjustments, suboptimal symptom control, or TUE-related treatment '
        'compromises. Second, the review provides actionable, disease-specific guidance that S&C professionals '
        'can immediately apply when working with athletes on prohibited medications or navigating the TUE '
        'process. Third, the manuscript addresses an emerging need: as clinical guidelines increasingly '
        'recommend agents on the WADA Prohibited List (e.g., GLP-1 receptor agonists for type 2 diabetes), '
        'S&C professionals require a comprehensive reference to anticipate and manage these discrepancies.'
    ))

    # Author perspective and JADA anecdote (toned down, as illustrative)
    add_para(doc, (
        'A distinctive feature of this review is the integration of the author\u2019s perspective as both a '
        'clinician and a former competitive bodybuilder regulated under the Japan Anti-Doping Agency (JADA). '
        'This dual perspective informed the identification of a gap that is under-recognized in the academic '
        'literature: the variability in TUE access across jurisdictions and competitive tiers. As an '
        'illustrative anecdote, the author attended an official JADA anti-doping education session where it '
        'was explicitly stated that TUE submissions were limited to athletes with prior top placements at '
        'national competitions or prior anti-doping rule violations. While this may reflect a localized '
        'operational practice rather than formal JADA policy, it illustrates the broader structural concern '
        'that TUE access may not be uniformly available to all athletes subject to anti-doping regulations\u2014'
        'a disparity that, combined with limited access to physicians knowledgeable in both anti-doping rules '
        'and evidence-based pharmacotherapy, creates significant barriers to equitable care and may narrow the '
        'base of sport participation.'
    ))

    # Core argument
    add_para(doc, (
        'A central argument of this review is that standard clinical treatment should be the baseline '
        'entitlement for all athletes, with restrictions applied only where a specific substance confers a '
        'clear, demonstrated performance advantage beyond therapeutic restoration. The current system inverts '
        'this logic by making prohibition the default and requiring athletes to petition for exceptions. The '
        'treatment burden this imposes\u2014regulatory uncertainty, administrative complexity, limited access to '
        'knowledgeable physicians, and psychological distress\u2014represents a significant and under-recognized '
        'harm to athletes that this review aims to bring to wider attention.'
    ))

    # Closing
    add_para(doc, (
        'The manuscript has not been previously published and is not under consideration by any other journal. '
        'A version of the core argument was previously submitted to the British Journal of Sports Medicine as '
        'a Viewpoint (approximately 1,500 words); however, the present manuscript is a substantially '
        'different work\u2014a comprehensive scoping review (approximately 6,700 words) with systematic '
        'methodology, disease-specific analysis, and practical recommendations not present in the earlier '
        'submission. All authors have approved the manuscript and agree with its submission to the SCJ.'
    ))

    add_para(doc, (
        'We confirm that we have read and complied with the Instructions for Authors. The PRISMA-ScR '
        'checklist is included as a supplementary file.'
    ))

    add_para(doc, '')
    add_para(doc, 'Sincerely,')
    add_para(doc, '[Author Name], MD, CSCS\n[Institutional Affiliation]\nEmail: [email]\nTelephone: [phone]')

    path = os.path.join(OUT_DIR, 'SCJ_Cover_Letter_English.docx')
    doc.save(path)
    print(f'Saved: {path}')


def create_japanese():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15

    add_para(doc, '［日付］')
    add_para(doc, '')
    add_para(doc, '編集長\nStrength and Conditioning Journal\nNational Strength and Conditioning Association')
    add_para(doc, '')
    add_para(doc, '拝啓', bold=True)

    add_para(doc, (
        '「Mind the Gap: WADA治療使用特例規則と現行臨床診療ガイドラインの乖離に関するスコーピングレビュー」と題した'
        '原稿を、Strength and Conditioning Journalのナラティブレビュー論文としてご検討いただきたく、謹んで投稿いたします。'
    ))

    add_para(doc, (
        '本PRISMA-ScR準拠スコーピングレビューは、世界アンチ・ドーピング機構（WADA）の治療使用特例（TUE）規制枠組みと、'
        '7つの疾患領域（喘息、ADHD、2型糖尿病、男性性腺機能低下症、グルココルチコイド投与を要する疾患、心血管疾患、'
        '多嚢胞性卵巣症候群）における現行のエビデンスに基づく臨床診療ガイドラインとの乖離を体系的にマッピングしました。'
        '68件の関連文献を同定し、7領域すべてで臨床的に有意なギャップを確認しました。'
    ))

    add_para(doc, (
        '本原稿がSCJの読者層に特に適していると考える理由は以下の通りです。第一に、臨床─競技ギャップはS&C専門家が'
        '日常的に指導するアスリートに直接影響します。第二に、本レビューはS&C専門家が禁止薬物を服用中のアスリートや'
        'TUEプロセスのナビゲーションにおいて即座に活用できる、疾患別の実践的ガイダンスを提供します。第三に、臨床'
        'ガイドラインがWADA禁止表上の薬剤をますます推奨する中で、S&C専門家がこれらの乖離を予測・管理するための'
        '包括的な参考資料が必要とされています。'
    ))

    add_para(doc, (
        '本レビューの特徴的な点は、臨床医であると同時に日本アンチ・ドーピング機構（JADA）管轄下の元競技ボディビルダー'
        'としての著者の視点を統合していることです。この二重の視点は、学術文献で十分に認識されていないギャップの同定に'
        '寄与しました：管轄区域と競技レベル間のTUEアクセスの変動性です。例示的な逸話として、著者が参加した公式JADA'
        'アンチ・ドーピング講習会では、TUE提出は全国大会で上位入賞した選手か、過去に規則違反をした選手のみに限られると'
        '明示されました。これはJADAの正式な方針ではなく局所的な運用実態を反映している可能性がありますが、アンチ・'
        'ドーピング規則の対象となる全てのアスリートにTUEアクセスが均一に提供されていない可能性があるという、より広い'
        '構造的懸念を例証するものです。この格差は、アンチ・ドーピング規則とエビデンスに基づく薬物療法の両方に精通した'
        '医師へのアクセスの限界と相まって、公平なケアへの重大な障壁を生み出し、競技の裾野を狭める可能性があります。'
    ))

    add_para(doc, (
        '本レビューの中心的主張は、標準的臨床治療が全てのアスリートにとって基本的権利であるべきであり、制限は特定の'
        '物質が治療的回復を超える明確なパフォーマンス向上効果をもたらす場合にのみ適用されるべきであるということです。'
        '現行制度はこの論理を反転させ、禁止をデフォルトとしてアスリートに例外を請願させています。これが課す治療負担'
        '──規制の不確実性、行政的複雑さ、知識ある医師へのアクセス制限、心理的苦痛──は、本レビューがより広く注目を'
        '集めることを目指す、アスリートに対する重大かつ十分に認識されていない害を表しています。'
    ))

    add_para(doc, (
        '本原稿は未発表であり、他のジャーナルで審査中ではありません。核心的議論の一部は以前British Journal of Sports '
        'MedicineにViewpoint（約1,500語）として投稿されましたが、本原稿は体系的方法論、疾患別分析、実践的提言を含む'
        '包括的スコーピングレビュー（約6,700語）であり、先の投稿とは実質的に異なる著作です。全著者が原稿を承認し、'
        'SCJへの投稿に同意しています。'
    ))

    add_para(doc, '投稿規定を確認し遵守していることを確認いたします。PRISMA-ScRチェックリストは補足ファイルとして添付しています。')

    add_para(doc, '')
    add_para(doc, '敬具')
    add_para(doc, '[著者名], MD, CSCS\n[所属機関]\nEmail: [email]\nTelephone: [phone]')

    path = os.path.join(OUT_DIR, 'SCJ_Cover_Letter_Japanese.docx')
    doc.save(path)
    print(f'Saved: {path}')


if __name__ == '__main__':
    create_english()
    create_japanese()
