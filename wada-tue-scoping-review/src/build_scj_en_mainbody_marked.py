#!/usr/bin/env python3
"""Build a marked (track-changes style) main-body .docx against the uploaded base version."""
import difflib
import os
import re
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from create_scj_review_part1 import setup_styles, add_heading_styled
from scj_en_content import (
    ABSTRACT, INTRO, METHODS, RESULTS_CONTEXT, RESULTS_OVERVIEW, DISEASE_RESULTS,
    CROSS_CUTTING, CROSS_CUTTING_2, DISCUSSION, RECS, PRACTICAL_ITEMS,
    FIGURE_LEGENDS, REFS, DISCLOSURES,
)

BLUE = RGBColor(0x00, 0x00, 0xFF)
RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)


def _collect_headings():
    """Return a set of expected heading strings from the new content."""
    headings = {
        "ABSTRACT", "INTRODUCTION", "METHODS", "RESULTS", "DISCUSSION",
        "Recommendations for Harmonization", "PRACTICAL APPLICATIONS", "DISCLOSURES",
        "REFERENCES", "FIGURE LEGENDS",
    }
    for title in METHODS.keys():
        headings.add(title)
    for title in DISEASE_RESULTS.keys():
        headings.add(title)
    headings.add("Population and Epidemiological Context")
    headings.add("Overview of Clinical-Competition Gaps")
    headings.add("Cross-Cutting Analysis: Structural Drivers of Divergence")
    return headings


HEADINGS = _collect_headings()
LEVEL2 = {
    "Protocol and Registration", "Eligibility Criteria",
    "Information Sources and Search Strategy", "Selection of Sources of Evidence",
    "Data Charting Process", "Critical Appraisal of Individual Sources",
    "Synthesis of Results", "Population and Epidemiological Context",
    "Overview of Clinical-Competition Gaps", *METHODS.keys(), *DISEASE_RESULTS.keys(),
    "Cross-Cutting Analysis: Structural Drivers of Divergence",
    "Recommendations for Harmonization",
}


def _add_colored_runs(p, text, color, strike=False, size=12, bold=False, italic=False):
    """Add text runs with given color/strikethrough and basic formatting."""
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.strike = strike
    if bold:
        r.font.bold = True
    if italic:
        r.font.italic = True
    return r


def _tokenize(text):
    # Keep each word together with its trailing whitespace so runs remain readable
    return re.findall(r"\S+\s*", text)


def _word_diff_runs(old_text, new_text, size=12, bold=False):
    """Return list of (text, color, strike) tuples for a word-level diff."""
    old_tokens = _tokenize(old_text)
    new_tokens = _tokenize(new_text)
    runs = []
    sm = difflib.SequenceMatcher(None, old_tokens, new_tokens)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            runs.append(("".join(new_tokens[j1:j2]), BLACK, False))
        elif tag == "delete":
            runs.append(("".join(old_tokens[i1:i2]), BLUE, True))
        elif tag == "insert":
            runs.append(("".join(new_tokens[j1:j2]), RED, False))
        elif tag == "replace":
            # Show deleted old tokens then inserted new tokens
            runs.append(("".join(old_tokens[i1:i2]), BLUE, True))
            runs.append(("".join(new_tokens[j1:j2]), RED, False))
    return runs


def _add_runs_to_paragraph(p, runs, size=12, bold=False, italic=False, center=False):
    for text, color, strike in runs:
        if not text:
            continue
        _add_colored_runs(p, text, color, strike=strike, size=size, bold=bold, italic=italic)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_paragraph(doc, text, color=BLACK, strike=False, size=12, bold=False, first_indent=True, center=False):
    p = doc.add_paragraph()
    if first_indent and not bold and not center:
        p.paragraph_format.first_line_indent = Cm(1.27)
    _add_colored_runs(p, text, color, strike=strike, size=size, bold=bold, italic=center)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _para_style(text):
    if text in HEADINGS:
        size = 12 if text in LEVEL2 else 14
        return True, size
    return False, 12


def _extract_mainbody_paras(doc_path):
    doc = Document(doc_path)
    texts = [p.text for p in doc.paragraphs]
    # Find ABSTRACT heading in base mainbody
    start = 0
    for i, t in enumerate(texts):
        if t and t.strip() == "ABSTRACT":
            start = i
            break
    return [t for t in texts[start:] if t is not None]


def build_marked_mainbody(base_path, new_path, out_path):
    base_paras = _extract_mainbody_paras(base_path)
    new_paras = _extract_mainbody_paras(new_path)

    doc = Document()
    setup_styles(doc)

    sm = difflib.SequenceMatcher(None, base_paras, new_paras)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for t in new_paras[j1:j2]:
                is_heading, size = _para_style(t)
                _add_paragraph(doc, t, color=BLACK, size=size, bold=is_heading,
                               first_indent=not is_heading, center=False)
        elif tag == "delete":
            for t in base_paras[i1:i2]:
                is_heading, size = _para_style(t)
                _add_paragraph(doc, t, color=BLUE, strike=True, size=size, bold=is_heading,
                               first_indent=not is_heading, center=False)
        elif tag == "insert":
            for t in new_paras[j1:j2]:
                is_heading, size = _para_style(t)
                _add_paragraph(doc, t, color=RED, size=size, bold=is_heading,
                               first_indent=not is_heading, center=False)
        elif tag == "replace":
            base_block = base_paras[i1:i2]
            new_block = new_paras[j1:j2]
            if len(base_block) == len(new_block):
                for b, n in zip(base_block, new_block):
                    is_heading_b, size_b = _para_style(b)
                    is_heading_n, size_n = _para_style(n)
                    # Use the new paragraph's style; if heading level differs, prefer new
                    size = size_n
                    bold = is_heading_n
                    p = doc.add_paragraph()
                    if not bold:
                        p.paragraph_format.first_line_indent = Cm(1.27)
                    runs = _word_diff_runs(b, n, size=size, bold=bold)
                    _add_runs_to_paragraph(p, runs, size=size, bold=bold)
            else:
                # Block replaced by different number of paragraphs: show old deleted then new inserted
                for b in base_block:
                    is_heading, size = _para_style(b)
                    _add_paragraph(doc, b, color=BLUE, strike=True, size=size, bold=is_heading,
                                   first_indent=not is_heading, center=False)
                for n in new_block:
                    is_heading, size = _para_style(n)
                    _add_paragraph(doc, n, color=RED, size=size, bold=is_heading,
                                   first_indent=not is_heading, center=False)

    doc.save(out_path)
    print(f"Saved marked mainbody: {out_path}")


def main():
    base = "/home/ubuntu/attachments/542f4773-4dec-4359-9fd3-7354d542d44a/SCJ_Narrative_Review_mainbody.docx"
    new = os.path.join(os.path.dirname(os.path.dirname(__file__)), "manuscripts", "SCJ_Narrative_Review_mainbody.docx")
    out = os.path.join(os.path.dirname(os.path.dirname(__file__)), "manuscripts", "SCJ_Narrative_Review_mainbody_marked.docx")
    build_marked_mainbody(base, new, out)


if __name__ == "__main__":
    main()
