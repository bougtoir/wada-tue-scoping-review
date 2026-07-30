#!/usr/bin/env python3
"""Re-order SCJ mainbody references alphabetically, renumber citations, add figure call-outs."""
import difflib
import json
import re
import sys
import os

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from create_scj_review_part1 import setup_styles

BLUE = RGBColor(0x00, 0x00, 0xFF)
RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

CITE_RE = re.compile(r'(\(\s*\d{1,2}(?:\s*[-–—]\s*\d{1,2})?(?:\s*,\s*\d{1,2}(?:\s*[-–—]\s*\d{1,2})?)*\s*\))')
NUM_RE = re.compile(r'\d+')

HEADINGS_L1 = {"ABSTRACT", "INTRODUCTION", "METHODS", "RESULTS", "DISCUSSION", "REFERENCES", "FIGURE LEGENDS", "PRACTICAL APPLICATIONS", "DISCLOSURES"}
HEADINGS_L2 = {
    "Protocol and Registration",
    "Eligibility Criteria",
    "Information Sources and Search Strategy",
    "Selection of Sources of Evidence",
    "Data Charting Process",
    "Critical Appraisal of Individual Sources",
    "Synthesis of Results",
    "Population and Epidemiological Context",
    "Overview of Clinical-Competition Gaps",
    "Asthma and Beta-2 Agonists",
    "Attention-Deficit/Hyperactivity Disorder (ADHD)",
    "Type 2 Diabetes Mellitus and GLP-1 Receptor Agonists",
    "Male Hypogonadism and Testosterone Replacement",
    "Glucocorticoid Administration",
    "Cardiovascular Disease: Diuretics and Beta-Blockers",
    "PMOS and Female Fertility Treatment",
    "Cross-Cutting Analysis: Structural Drivers of Divergence",
    "Recommendations for Harmonization",
}

def _ref_sort_key(text):
    s = text.strip()
    if s.lower().startswith('the '):
        s = s[4:]
    # first two alphabetic words for surname handling
    words = re.findall(r"[A-Za-z'\u00C0-\u024F]+", s)
    return ' '.join(words[:2]).lower()

def _heading_level(text):
    t = text.strip()
    if t in HEADINGS_L1:
        return 1
    if t in HEADINGS_L2:
        return 2
    # heading if all uppercase and short
    if t.isupper() and len(t) < 80:
        return 1
    return 0

def _tokenize(text):
    return re.findall(r"\S+\s*", text)

def _map_citations(text, mapping):
    def repl(m):
        inner = m.group(0)[1:-1]
        parts = re.split(r'(\s*[-–—,]\s*)', inner)
        out = []
        for p in parts:
            n = p.strip()
            if NUM_RE.fullmatch(n):
                out.append(str(mapping.get(int(n), n)))
            else:
                out.append(p)
        return '(' + ''.join(out) + ')'
    return CITE_RE.sub(repl, text)

def _build_refs(ref_paras):
    entries = []
    for p in ref_paras:
        t = p['text'].strip()
        m = re.match(r'(\d+)\.\s*(.*)', t)
        if m:
            entries.append((int(m.group(1)), m.group(2)))
    # sort alphabetically (stable)
    sorted_entries = sorted(entries, key=lambda x: (_ref_sort_key(x[1]), x[1].lower()))
    old_to_new = {old: i+1 for i, (old, _) in enumerate(sorted_entries)}
    new_refs = [f"{i+1}. {text}" for i, (_, text) in enumerate(sorted_entries)]
    return old_to_new, new_refs

def _add_runs_colored(p, text, color, strike=False, size=12, bold=False, italic=False, superscript=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = color
    r.font.strike = strike
    r.font.bold = bold
    r.font.italic = italic
    r.font.superscript = superscript
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return r

def _add_citation_runs(p, text, color=BLACK, strike=False, size=12, bold=False, italic=False):
    """Render text with citation groups as font-based superscript, preserving color/strike."""
    for part in CITE_RE.split(text):
        if CITE_RE.fullmatch(part):
            _add_runs_colored(p, '(', color, strike=strike, size=size, bold=bold, italic=italic)
            inner = part[1:-1]
            for token in re.findall(r'\d+|\s*[-–—,]\s*', inner):
                n = token.strip()
                if NUM_RE.fullmatch(n):
                    _add_runs_colored(p, n, color, strike=strike, size=size, bold=bold, italic=italic, superscript=True)
                else:
                    _add_runs_colored(p, token, color, strike=strike, size=size, bold=bold, italic=italic)
            _add_runs_colored(p, ')', color, strike=strike, size=size, bold=bold, italic=italic)
        else:
            _add_runs_colored(p, part, color, strike=strike, size=size, bold=bold, italic=italic)

def _add_paragraph_clean(doc, text, level=0, figure_callout=False):
    p = doc.add_paragraph()
    if figure_callout:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_citation_runs(p, text, italic=True, size=12)
        return p
    if level == 1:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.first_line_indent = Cm(1.27)
        _add_citation_runs(p, text, size=12)
    return p

def _add_ref_paragraph_clean(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 2.0

def _word_diff_runs(old_text, new_text, size=12, bold=False):
    old_tokens = _tokenize(old_text)
    new_tokens = _tokenize(new_text)
    runs = []
    sm = difflib.SequenceMatcher(None, old_tokens, new_tokens)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            runs.append((''.join(new_tokens[j1:j2]), BLACK, False))
        elif tag == 'delete':
            runs.append((''.join(old_tokens[i1:i2]), BLUE, True))
        elif tag == 'insert':
            runs.append((''.join(new_tokens[j1:j2]), RED, False))
        elif tag == 'replace':
            runs.append((''.join(old_tokens[i1:i2]), BLUE, True))
            runs.append((''.join(new_tokens[j1:j2]), RED, False))
    return runs

def _add_paragraph_marked(doc, text, size=12, level=0, is_insert=False, is_delete=False):
    p = doc.add_paragraph()
    if level == 1:
        size = 14
        bold = True
    elif level == 2:
        size = 12
        bold = True
    else:
        bold = False
    # Numbered list items (references, recommendations) should not have issue numbers superscripted
    is_numbered_item = re.match(r'^\d+\.\s', text.strip()) is not None
    color = RED if is_insert else (BLUE if is_delete else BLACK)
    if is_insert or is_delete or is_numbered_item:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = color
        run.font.strike = is_delete
        run.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    else:
        if level == 0:
            p.paragraph_format.first_line_indent = Cm(1.27)
        _add_citation_runs(p, text, size=size, bold=bold)
    if level:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    return p

def _build_marked_doc(old_texts, new_texts, out_path):
    doc = Document()
    setup_styles(doc)
    sm = difflib.SequenceMatcher(None, old_texts, new_texts)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == 'equal':
            for t in new_texts[j1:j2]:
                lvl = _heading_level(t)
                _add_paragraph_marked(doc, t, level=lvl)
        elif tag == 'delete':
            for t in old_texts[i1:i2]:
                lvl = _heading_level(t)
                _add_paragraph_marked(doc, t, level=lvl, is_delete=True)
        elif tag == 'insert':
            for t in new_texts[j1:j2]:
                lvl = _heading_level(t)
                _add_paragraph_marked(doc, t, level=lvl, is_insert=True)
        elif tag == 'replace':
            old_block = old_texts[i1:i2]
            new_block = new_texts[j1:j2]
            for old_t, new_t in zip(old_block, new_block):
                # Numbered paragraphs that have been reordered (e.g., references, recommendations)
                # should be shown as whole-paragraph delete + insert, not inline word diff.
                if re.match(r'^\d+\.\s', old_t.strip()) and re.match(r'^\d+\.\s', new_t.strip()):
                    _add_paragraph_marked(doc, old_t, level=_heading_level(old_t), is_delete=True)
                    _add_paragraph_marked(doc, new_t, level=_heading_level(new_t), is_insert=True)
                elif difflib.SequenceMatcher(None, old_t, new_t).quick_ratio() < 0.5:
                    _add_paragraph_marked(doc, old_t, level=_heading_level(old_t), is_delete=True)
                    _add_paragraph_marked(doc, new_t, level=_heading_level(new_t), is_insert=True)
                else:
                    lvl = _heading_level(new_t)
                    p = doc.add_paragraph()
                    if lvl:
                        size = 14 if lvl == 1 else 12
                        bold = True
                    else:
                        size = 12
                        bold = False
                        p.paragraph_format.first_line_indent = Cm(1.27)
                    for txt, color, strike in _word_diff_runs(old_t, new_t, size=size, bold=bold):
                        if not txt:
                            continue
                        _add_citation_runs(p, txt, color=color, strike=strike, size=size, bold=bold)
                    if lvl:
                        p.paragraph_format.space_before = Pt(12)
                        p.paragraph_format.space_after = Pt(6)
            # Any leftover paragraphs in old block are deletions, leftovers in new block are insertions
            if len(old_block) > len(new_block):
                for t in old_block[len(new_block):]:
                    _add_paragraph_marked(doc, t, level=_heading_level(t), is_delete=True)
            elif len(new_block) > len(old_block):
                for t in new_block[len(old_block):]:
                    _add_paragraph_marked(doc, t, level=_heading_level(t), is_insert=True)
    doc.save(out_path)

def main():
    import zipfile
    from lxml import etree
    NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    DEL_TAGS = {f'{{{NS["w"]}}}del', f'{{{NS["w"]}}}moveFrom'}

    def is_deleted(elem):
        for anc in elem.iterancestors():
            if anc.tag in DEL_TAGS:
                return True
        return False

    def para_text(p):
        parts = []
        for t in p.iter():
            if t.tag == f'{{{NS["w"]}}}t' or t.tag == f'{{{NS["w"]}}}delText':
                if not is_deleted(t):
                    parts.append(t.text or '')
        return ''.join(parts)

    def para_style(p):
        pPr = p.find('w:pPr', NS)
        if pPr is not None:
            ps = pPr.find('w:pStyle', NS)
            if ps is not None:
                return ps.get(f'{{{NS["w"]}}}val')
        return None

    docx_path = '/home/ubuntu/attachments/96ef9444-26fc-45be-88ba-89e6e64a0f42/SCJ_1st_mainbody_trackchange.docx'
    out_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'manuscripts')
    os.makedirs(out_dir, exist_ok=True)

    with zipfile.ZipFile(docx_path) as z:
        xml = z.read('word/document.xml')
    root = etree.fromstring(xml)
    body = root.find('.//w:body', NS)
    paras = []
    for i, p in enumerate(body.findall('.//w:p', NS)):
        t = para_text(p)
        if t.strip():
            paras.append({'idx': i, 'style': para_style(p), 'text': t})

    # find references and figure legends
    ref_idx = None
    fig_idx = None
    for i, p in enumerate(paras):
        t = p['text'].strip()
        if t == 'REFERENCES':
            ref_idx = i
        if t == 'FIGURE LEGENDS':
            fig_idx = i
            break

    old_to_new, new_refs = _build_refs(paras[ref_idx+1:fig_idx])

    # Build new paragraphs: body with renumbered citations, then new references, then figure legends
    new_paras_texts = []
    # body before references (renumber citations + add call-outs)
    body_texts = [_map_citations(p['text'], old_to_new) for p in paras[:ref_idx+1]]

    # Add figure call-outs
    for fig in range(1, 6):
        inserted = False
        pattern = re.compile(rf'Figure\s+{fig}(?!\d)')
        for i, t in enumerate(body_texts):
            if pattern.search(t):
                body_texts.insert(i+1, f'[Insert Figure {fig} here]')
                inserted = True
                break
        if not inserted:
            print(f'Warning: Figure {fig} first mention not found')

    new_paras_texts.extend(body_texts)
    new_paras_texts.extend(new_refs)
    # figure legends onwards (renumber? no)
    new_paras_texts.extend([p['text'] for p in paras[fig_idx:]])

    # Build clean final docx
    clean_path = os.path.join(out_dir, 'SCJ_Narrative_Review_mainbody.docx')
    clean_doc = Document()
    setup_styles(clean_doc)
    for t in new_paras_texts:
        lvl = _heading_level(t)
        fig_callout = t.startswith('[Insert Figure') and t.endswith('here]')
        if fig_callout:
            _add_paragraph_clean(clean_doc, t, figure_callout=True)
        elif t.startswith('REFERENCES'):
            _add_paragraph_clean(clean_doc, t, level=lvl)
        elif ref_idx is not None and t in new_refs:
            _add_ref_paragraph_clean(clean_doc, t)
        else:
            _add_paragraph_clean(clean_doc, t, level=lvl)
    clean_doc.save(clean_path)
    print('Clean final saved:', clean_path)

    # Build marked docx
    old_texts = [p['text'] for p in paras]
    marked_path = os.path.join(out_dir, 'SCJ_Narrative_Review_mainbody_marked.docx')
    _build_marked_doc(old_texts, new_paras_texts, marked_path)
    print('Marked doc saved:', marked_path)

    # Save mapping for response update
    mapping_path = os.path.join(out_dir, 'citation_mapping.json')
    with open(mapping_path, 'w') as f:
        json.dump({str(k): v for k, v in old_to_new.items()}, f)
    print('Mapping saved:', mapping_path)

if __name__ == '__main__':
    main()
