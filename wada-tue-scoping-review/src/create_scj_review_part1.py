#!/usr/bin/env python3
"""Create SCJ Full Narrative Review .docx - Part 1: Setup and helper functions."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_DIR = os.path.join(BASE_DIR, 'manuscripts')
FIG_DIR = os.path.join(BASE_DIR, 'figures')
os.makedirs(OUTPUT_DIR, exist_ok=True)

def _add_page_number_footer(section):
    """Add a centered PAGE field to the section footer."""
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    fld_separate = OxmlElement('w:fldChar')
    fld_separate.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def setup_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = 2.0
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        _add_page_number_footer(section)

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    else:
        run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

CITE_RE = re.compile(r'(\(\s*\d{1,3}(?:\s*(?:,|-|–|—)\s*\d{1,3})*\s*\))')
NUM_RE = re.compile(r'\d+')


def _add_runs(p, text, font_size, **kwargs):
    """Add text to paragraph, rendering citation numbers as font-based superscript."""
    def _run(s):
        r = p.add_run(s)
        r.font.size = Pt(font_size)
        for k, v in kwargs.items():
            setattr(r.font, k, v)
        return r

    for part in CITE_RE.split(text):
        if CITE_RE.fullmatch(part):
            # Citation group: normal parentheses, superscript numbers
            _run('(')
            inner = part[1:-1]
            for token in re.findall(r'\d+|\s*[,–—-]\s*', inner):
                if NUM_RE.fullmatch(token.strip()):
                    r = _run(token.strip())
                    r.font.superscript = True
                else:
                    _run(token)
            _run(')')
        else:
            _run(part)


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    _add_runs(p, text, 12)
    return p

def add_body_ni(doc, text):
    p = doc.add_paragraph()
    _add_runs(p, text, 12)
    return p

def add_figure(doc, filename, caption):
    fig_path = os.path.join(FIG_DIR, filename)
    if os.path.exists(fig_path):
        doc.add_picture(fig_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p = doc.add_paragraph()
        run = p.add_run(f'[Figure placeholder: {filename}]')
        run.font.size = Pt(10); run.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_figure_legend(doc, caption)


def add_figure_legend(doc, caption):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.line_spacing = 1.5
    _add_runs(p, caption, 10, italic=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def add_ref(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 2.0

print("Part 1 helpers loaded.")
