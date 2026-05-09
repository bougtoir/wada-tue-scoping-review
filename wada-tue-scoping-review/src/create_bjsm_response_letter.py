#!/usr/bin/env python3
"""Create BJSM Response Letter (EN + JP) documenting rebuttal to reviewer comments."""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "manuscripts")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def setup_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.5
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def add_para(doc, text, bold=False, italic=False, size=Pt(11), indent=False,
             space_before=Pt(0), space_after=Pt(4), alignment=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = size
    run.bold = bold
    run.italic = italic
    return p


def add_superscript_para(doc, text, bold=False, indent=False):
    """Add paragraph with {n} markers converted to superscript."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1.27)
    parts = re.split(r'(\{[^}]+\})', text)
    for part in parts:
        if part.startswith('{') and part.endswith('}'):
            run = p.add_run(part[1:-1])
            run.font.size = Pt(11)
            run.font.superscript = True
        else:
            run = p.add_run(part)
            run.font.size = Pt(11)
            run.bold = bold
    return p


# ============================================================
# ENGLISH VERSION
# ============================================================
def create_english():
    doc = Document()
    setup_styles(doc)

    add_para(doc, 'Response to BJSM Decision and Reviewer Comments',
             bold=True, size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=Pt(12))

    add_para(doc, 'Manuscript: Mind the Gap: Clinical Practice Guidelines and WADA Therapeutic Use '
             'Exemptions Are Diverging\u2014A Call for Harmonization',
             italic=True, space_after=Pt(12))

    add_para(doc, 'Date: May 2026', space_after=Pt(12))

    # ---- Editor response ----
    add_para(doc, 'Dear Associate Editor,', space_after=Pt(8))

    add_para(doc, 'Thank you for the opportunity to submit our viewpoint to the BJSM and for the '
             'detailed feedback from the reviewers. While we respectfully accept the editorial '
             'decision, we wish to clarify several points where we believe the reviewers may have '
             'misinterpreted our arguments. We document these clarifications below for the record '
             'and to inform our revised submission to the Strength and Conditioning Journal (SCJ), '
             'where we will present the full scoping review with expanded practical applications.')

    add_para(doc, 'We emphasize that our manuscript was submitted as a Viewpoint/Commentary '
             '(approximately 1,500 words), which necessarily limited the depth of discussion on '
             'each topic. The full scoping review, now being prepared for SCJ, addresses many of '
             'the reviewers\u2019 concerns in greater detail.',
             space_after=Pt(12))

    # ---- REVIEWER 1 ----
    add_para(doc, 'Response to Reviewer 1', bold=True, size=Pt(12), space_before=Pt(12))

    add_para(doc, 'Reviewer 1 Comment:', bold=True, space_before=Pt(8))
    add_para(doc, '"The author does not demonstrate an understanding of anti-doping processes '
             'including the TUE process. Athletes are not required to cease prohibited medications, '
             'TUEs permit them to continue them and participate in sport."',
             italic=True, indent=True)

    add_para(doc, 'Response:', bold=True, space_before=Pt(8))
    add_para(doc, 'We respectfully disagree with this characterization. The lead author is a '
             'physician (MD) with CSCS certification and personal experience as a competitive '
             'athlete subject to anti-doping testing under the Japan Anti-Doping Agency (JADA). '
             'We are fully aware that the TUE system exists precisely to permit athletes with '
             'legitimate medical conditions to use otherwise prohibited substances.')

    add_para(doc, 'Our argument is not that TUEs do not exist, but rather that structural gaps '
             'between the TUE regulatory framework and current clinical practice guidelines create '
             'practical barriers that compromise the system\u2019s intended function. Specifically:')

    items_r1 = [
        ('TUE access is not universal or equitable. ',
         'In the author\u2019s direct experience as a competitive bodybuilder under JADA jurisdiction, '
         'athletes were informed during official anti-doping education sessions that TUE applications '
         'could only be submitted by athletes who had previously placed in top positions at national '
         'competitions or who had a prior anti-doping rule violation. This effectively excludes '
         'emerging athletes, newcomers, and those competing at sub-elite levels from TUE access, '
         'creating a two-tier system where the right to receive evidence-based medical treatment '
         'while competing depends on competitive achievement rather than medical need.'),
        ('The gap between clinical standard of care and TUE criteria is the core issue. ',
         'Even when TUE applications are theoretically available, the criteria for approval may '
         'diverge from clinical guidelines. For example, WADA\u2019s restriction of testosterone TUEs '
         'to organic etiologies excludes functional hypogonadism, which the Endocrine Society and '
         'EAU guidelines recognize as a treatable condition. The TUE system cannot bridge a gap '
         'when the eligibility criteria themselves are misaligned with clinical evidence.'),
        ('The viewpoint format constrained our discussion. ',
         'The BJSM viewpoint format (~1,500 words) required us to present a high-level overview '
         'across seven disease areas. The full scoping review (now being prepared for SCJ at '
         '5,000\u20138,000 words) provides detailed disease-specific analysis including the TUE process, '
         'approval criteria, and practical implications.'),
    ]
    for bold_part, rest in items_r1:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-0.63)
        run = p.add_run('\u2022 ')
        run.font.size = Pt(11)
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(rest)
        run.font.size = Pt(11)

    add_para(doc, 'Reviewer 1\u2019s comment that "the WADA prohibited list and TUE process is not just '
             'about the welfare of the athletes but also protection of fairness in sport" is a point '
             'we explicitly address in our manuscript. Our scoping review identifies the tension '
             'between these dual objectives as a structural driver of the clinical-competition gap.',
             space_before=Pt(8))

    # ---- REVIEWER 2 ----
    add_para(doc, 'Response to Reviewer 2', bold=True, size=Pt(12), space_before=Pt(16))

    # Point 1: Hypogonadism
    add_para(doc, 'Point 1 \u2013 Male Hypogonadism and Low Testosterone:', bold=True, space_before=Pt(8))
    add_para(doc, 'Reviewer comment: Non-organic causes are theoretically reversible; aging males '
             'with borderline T from non-organic cause "will NEVER get a TUE"; T is a "VERY potent '
             'performance enhancer"; TUE requirements should only apply to elite athletes.',
             italic=True, indent=True)

    add_para(doc, 'Response:', bold=True, space_before=Pt(4))
    add_para(doc, 'We agree that testosterone is a potent performance enhancer and that this fact '
             'appropriately informs anti-doping policy. However, the reviewer\u2019s statement that an '
             'aging male with non-organic hypogonadism "will NEVER get a TUE" precisely illustrates '
             'our argument: when clinical guidelines (Endocrine Society 2018, EAU 2024) recommend '
             'treatment and WADA policy categorically excludes it, a clinical-competition gap exists '
             'by definition. We do not argue that all hypogonadal athletes should receive TUEs; '
             'we argue that the categorical exclusion of functional etiologies deserves re-evaluation '
             'in light of evolving clinical evidence, including the TRAVERSE trial\u2019s cardiovascular '
             'safety data.')

    add_para(doc, 'The reviewer\u2019s observation that "they are rarely elite" and that "TUE requirements '
             'should only be applied to elite athletes and keep out of recreational sport" supports '
             'our recommendation for a tiered approach, which we develop further in the SCJ manuscript.')

    # Point 2: Asthma
    add_para(doc, 'Point 2 \u2013 Asthma Treatment:', bold=True, space_before=Pt(8))
    add_para(doc, 'Reviewer comment: Biggest risk is with salbutamol, which is not mentioned. '
             'Athletes can apply for TUE if over permitted dose.',
             italic=True, indent=True)

    add_para(doc, 'Response:', bold=True, space_before=Pt(4))
    add_para(doc, 'We appreciate this feedback and accept that the salbutamol/albuterol threshold risk '
             'deserves more prominent discussion. In our viewpoint, we did address the concern that '
             'WADA\u2019s permissive approach to inhaled salbutamol without TUE inadvertently incentivizes '
             'SABA monotherapy\u2014which GINA 2025 explicitly recommends against. However, we acknowledge '
             'that the specific threshold exceedance risk with salbutamol (1600 \u03bcg/24h, 800 \u03bcg/8h '
             'limits and the urine decision limit of 1000 ng/mL) was insufficiently discussed. '
             'The full SCJ manuscript now includes expanded analysis of both salbutamol and formoterol '
             'threshold risks.')

    # Point 3: ADHD
    add_para(doc, 'Point 3 \u2013 ADHD:', bold=True, space_before=Pt(8))
    add_para(doc, 'Reviewer comment: Athletes can apply for TUE and use stimulants during competition '
             'if required for regular treatment.',
             italic=True, indent=True)

    add_para(doc, 'Response:', bold=True, space_before=Pt(4))
    add_para(doc, 'We understand that TUE applications for stimulant medications in ADHD are possible, '
             'and the revised SCJ manuscript explicitly discusses this pathway. However, TUE availability '
             'in theory does not equate to equitable access in practice. As noted above, the author\u2019s '
             'personal experience in JADA-regulated competitive bodybuilding revealed that TUE submission '
             'was restricted to athletes with prior top placements or rule violations. This operational '
             'restriction\u2014whether reflecting JADA-specific policy, testing capacity limitations, or '
             'sport-specific implementation\u2014means that emerging or sub-elite athletes diagnosed with '
             'ADHD may have no practical mechanism to obtain a TUE, even when clinical guidelines '
             'unequivocally recommend stimulant pharmacotherapy as first-line treatment.')

    add_para(doc, 'This jurisdictional variability in TUE access is precisely the type of systemic '
             'issue our scoping review aims to highlight.')

    # Point 4: GLP-1
    add_para(doc, 'Point 4 \u2013 GLP-1 Receptor Agonists:', bold=True, space_before=Pt(8))
    add_para(doc, 'Reviewer comment: Currently permitted; on monitoring list; if prohibited, TUE '
             'option would exist for medical indications like type 2 diabetes.',
             italic=True, indent=True)

    add_para(doc, 'Response:', bold=True, space_before=Pt(4))
    add_para(doc, 'We agree with the reviewer that GLP-1 receptor agonists are currently permitted '
             'and that TUE pathways would likely be established if they are prohibited. Our manuscript '
             'explicitly acknowledges this. Our concern is prospective: given the trajectory toward '
             'prohibition (monitoring program inclusion since 2024), we advocate for proactive impact '
             'assessment before prohibition rather than reactive TUE framework development afterward. '
             'The reviewer\u2019s suggestion that prohibition would likely be limited to weight-category '
             'and aesthetic sports is a nuanced point that we incorporate in the SCJ revision.')

    # Point 5: TUE guideline review
    add_para(doc, 'Point 5 \u2013 TUE Guideline Review Process:', bold=True, space_before=Pt(8))
    add_para(doc, 'Reviewer comment: Formal review process exists; most guidelines updated in last '
             '2 years; one outstanding review (sleep disorders).',
             italic=True, indent=True)

    add_para(doc, 'Response:', bold=True, space_before=Pt(4))
    add_para(doc, 'We appreciate this clarification and acknowledge that WADA has made progress in '
             'updating TUE Physician Guidelines. Our recommendation for a formal linkage mechanism '
             'between clinical guideline revisions and TUE guideline review is intended to complement, '
             'not replace, the existing process. The temporal mismatch we identified refers to the '
             'structural lag that can occur between major clinical guideline revisions and corresponding '
             'TUE guideline updates, even with an active review process. The SCJ manuscript now includes '
             'a more balanced acknowledgment of recent update efforts.')

    # ---- Concluding remarks ----
    add_para(doc, 'Concluding Remarks', bold=True, size=Pt(12), space_before=Pt(16))

    add_para(doc, 'We note that both reviewers appear to approach this topic primarily from a '
             'regulatory perspective, which is valuable but represents one side of a multifaceted issue. '
             'Our manuscript is written from the dual perspective of a physician-athlete who has '
             'personally experienced the anti-doping system as a competitive bodybuilder in Japan '
             'and who treats patients navigating these same challenges. This experiential perspective '
             'informs our identification of practical barriers that may not be apparent from the '
             'regulatory framework alone.')

    add_para(doc, 'We wish to emphasize a fundamental point: the existence of a TUE pathway does not, '
             'in itself, ensure equitable access to evidence-based medical care. The current system '
             'places the burden on athletes to navigate complex regulations, secure specialist '
             'documentation, and petition for exceptions to receive standard treatment. Few athletes '
             'have access to sports medicine physicians with expertise in both anti-doping regulations '
             'and their specific condition; most are treated by general practitioners or specialists '
             'unfamiliar with WADA rules. We argue that standard clinical treatment should be the '
             'baseline entitlement for all athletes, with restrictions applied only where a substance '
             'confers a clear performance advantage beyond therapeutic restoration. The treatment '
             'stress imposed on athletes\u2014regulatory uncertainty, administrative burden, and '
             'psychological distress from navigating these systems\u2014is itself a significant and '
             'under-recognized harm that our review aims to bring to wider attention.')

    add_para(doc, 'We will proceed with submission to the Strength and Conditioning Journal, where '
             'the full scoping review format (5,000\u20138,000 words) allows for the comprehensive '
             'treatment that both reviewers\u2019 comments suggest is needed. The SCJ manuscript '
             'incorporates the valid points raised by both reviewers, including expanded salbutamol '
             'analysis, more nuanced GLP-1 discussion, acknowledgment of recent TUE guideline '
             'updates, and\u2014critically\u2014a new section on jurisdictional variability in TUE access '
             'based on the author\u2019s first-hand experience.')

    add_para(doc, 'Respectfully,', space_before=Pt(16))
    add_para(doc, '[Author Name], MD, CSCS')

    output = os.path.join(OUTPUT_DIR, 'BJSM_Response_Letter_English.docx')
    doc.save(output)
    print(f"Saved: {output}")


# ============================================================
# JAPANESE VERSION
# ============================================================
def create_japanese():
    doc = Document()
    setup_styles(doc)

    add_para(doc, 'BJSM 編集判定およびレビュアーコメントへの回答',
             bold=True, size=Pt(14), alignment=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=Pt(12))

    add_para(doc, '原稿: Mind the Gap: Clinical Practice Guidelines and WADA Therapeutic Use '
             'Exemptions Are Diverging\u2014A Call for Harmonization',
             italic=True, space_after=Pt(12))

    add_para(doc, '日付: 2026年5月', space_after=Pt(12))

    add_para(doc, '編集委員長 殿', space_after=Pt(8))

    add_para(doc, 'このたびはBJSMへの投稿の機会をいただき、また査読者から詳細なフィードバックをいただきましたことに'
             '感謝申し上げます。編集判定を謹んで受け入れますが、査読者が我々の議論を誤解していると思われるいくつかの点に'
             'ついて明確化させていただきたく存じます。以下に記録として、また Strength and Conditioning Journal（SCJ）への'
             '改訂投稿に向けた情報として、これらの明確化を記載いたします。')

    add_para(doc, '本原稿はViewpoint/Commentary（約1,500語）として投稿したものであり、各トピックの議論の深さには'
             '必然的に制約がありました。現在SCJ向けに準備中の完全なスコーピングレビュー（5,000〜8,000語）では、'
             '査読者の懸念の多くをより詳細に扱っています。',
             space_after=Pt(12))

    # ---- REVIEWER 1 ----
    add_para(doc, 'レビュアー1への回答', bold=True, size=Pt(12), space_before=Pt(12))

    add_para(doc, 'レビュアー1のコメント:', bold=True, space_before=Pt(8))
    add_para(doc, '「著者はTUEプロセスを含むアンチ・ドーピングプロセスの理解を示していない。'
             'アスリートは禁止薬を中止する必要はなく、TUEにより継続使用と競技参加が認められている。」',
             italic=True, indent=True)

    add_para(doc, '回答:', bold=True, space_before=Pt(8))
    add_para(doc, 'この評価に対し、敬意をもって異議を唱えます。筆頭著者はCSCS認定を有する医師（MD）であり、'
             '日本アンチ・ドーピング機構（JADA）管轄下のアンチ・ドーピング検査を受ける競技選手としての実体験を有して'
             'います。TUE制度が、正当な医学的状態を有するアスリートに禁止物質の使用を許可するために存在していることは'
             '十分に理解しています。')

    add_para(doc, '我々の主張は「TUEが存在しない」ということではなく、TUE規制枠組みと現行臨床診療ガイドラインの間の'
             '構造的ギャップが、制度の意図する機能を損なう実践的障壁を生み出しているということです。具体的には：')

    items_r1_jp = [
        ('TUEアクセスは普遍的でも公平でもない。',
         '著者が競技ボディビルダーとしてJADA管轄下で活動した直接的経験において、公式アンチ・ドーピング講習会で'
         '「TUE申請は全国大会で上位入賞した選手か、過去にアンチ・ドーピング規則違反をした選手のみが提出できる」'
         'と説明されました。これは事実上、新興選手、初参加の選手、サブエリートレベルの選手をTUEアクセスから排除し、'
         'エビデンスに基づく医療を受けながら競技する権利が医学的必要性ではなく競技実績に依存する二層構造を生み出して'
         'います。'),
        ('臨床標準治療とTUE基準の間のギャップが核心的問題である。',
         'TUE申請が理論的に可能な場合でも、承認基準が臨床ガイドラインと乖離している場合があります。例えば、WADAの'
         'テストステロンTUEの器質的病因への限定は、内分泌学会（2018）やEAU（2024）ガイドラインが治療可能な状態として'
         '認める機能性性腺機能低下症を除外します。TUE制度は、適格基準自体が臨床エビデンスと不整合である場合、'
         'ギャップを橋渡しすることができません。'),
        ('Viewpoint形式が議論を制約した。',
         'BJSMのViewpoint形式（約1,500語）では、7つの疾患領域にわたる概要の提示が求められました。SCJ向けの完全な'
         'スコーピングレビュー（5,000〜8,000語）では、TUEプロセス、承認基準、実践的含意を含む詳細な疾患別分析を'
         '提供します。'),
    ]
    for bold_part, rest in items_r1_jp:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-0.63)
        run = p.add_run('\u2022 ')
        run.font.size = Pt(11)
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(rest)
        run.font.size = Pt(11)

    add_para(doc, 'レビュアー1の「WADA禁止表とTUEプロセスはアスリートの福祉だけでなく、スポーツの公正性の保護も'
             '目的としている」というコメントは、我々の原稿で明示的に扱っている点です。我々のスコーピングレビューは、'
             'これら二重の目的間の緊張を臨床─競技ギャップの構造的要因として同定しています。',
             space_before=Pt(8))

    # ---- REVIEWER 2 ----
    add_para(doc, 'レビュアー2への回答', bold=True, size=Pt(12), space_before=Pt(16))

    # Point 1
    add_para(doc, 'ポイント1 \u2013 男性性腺機能低下症と低テストステロン:', bold=True, space_before=Pt(8))
    add_para(doc, 'レビュアーコメント: 非器質的原因は理論的に「可逆的」である。テストステロンは「非常に強力な'
             'パフォーマンス向上物質」であり、加齢男性が非器質的原因でTUEを取得することは「絶対にない」。'
             'TUE要件はエリートアスリートにのみ適用し、レクリエーションスポーツからは排除すべき。',
             italic=True, indent=True)

    add_para(doc, '回答:', bold=True, space_before=Pt(4))
    add_para(doc, 'テストステロンが強力なパフォーマンス向上物質であること、この事実がアンチ・ドーピング政策に'
             '適切に反映されるべきであることに同意します。しかし、レビュアーが「加齢男性は非器質的原因では絶対に'
             'TUEを取得できない」と述べていることは、まさに我々の主張を例証しています：臨床ガイドライン（内分泌学会'
             '2018、EAU 2024）が治療を推奨し、WADA政策がそれをカテゴリカルに排除する場合、定義上、臨床─競技ギャップ'
             'が存在します。')

    add_para(doc, 'レビュアーの「彼らがエリートであることは稀」であり「TUE要件はエリートアスリートにのみ適用し'
             'レクリエーションスポーツからは排除すべき」という指摘は、SCJ原稿でさらに展開する段階的アプローチの'
             '提言を支持するものです。')

    # Point 2
    add_para(doc, 'ポイント2 \u2013 喘息治療:', bold=True, space_before=Pt(8))
    add_para(doc, 'レビュアーコメント: 最大のリスクはサルブタモールにある。許容用量を超えた場合、TUE申請可能。',
             italic=True, indent=True)

    add_para(doc, '回答:', bold=True, space_before=Pt(4))
    add_para(doc, 'このフィードバックに感謝し、サルブタモール/アルブテロールの閾値リスクのより詳細な議論が必要で'
             'あったことを認めます。Viewpointでは、WADAの吸入サルブタモールに対する寛容なアプローチがSABA単独療法への'
             '依存を意図せず促すことに言及しましたが、サルブタモールの具体的な閾値超過リスク（24時間1600 \u03bcg、8時間'
             '800 \u03bcg制限、尿中判定限界1000 ng/mL）の議論は不十分でした。SCJ完全原稿では、サルブタモールと'
             'ホルモテロール両方の閾値リスクの拡充された分析を含めています。')

    # Point 3
    add_para(doc, 'ポイント3 \u2013 ADHD:', bold=True, space_before=Pt(8))
    add_para(doc, 'レビュアーコメント: アスリートはTUEを申請し、定期治療が必要であれば競技中も使用可能。',
             italic=True, indent=True)

    add_para(doc, '回答:', bold=True, space_before=Pt(4))
    add_para(doc, 'ADHD刺激薬のTUE申請が可能であることは理解しており、改訂SCJ原稿ではこの経路を明示的に議論して'
             'います。しかし、理論上のTUE利用可能性は実践上の公平なアクセスを意味しません。上述のとおり、著者の'
             'JADA管轄下の競技ボディビルにおける個人的経験では、TUE提出は過去に上位入賞した選手または規則違反した'
             '選手に限定されていました。この運営上の制限は──それがJADA固有の方針、検査キャパシティの制約、'
             '競技固有の実施方法のいずれに起因するかにかかわらず──ADHDと診断された新興選手またはサブエリート選手が、'
             '臨床ガイドラインが第一選択治療として刺激薬を明確に推奨しているにもかかわらず、TUEを取得する実際的な'
             '手段を持たない可能性があることを意味します。')

    add_para(doc, 'このTUEアクセスの管轄間変動は、まさに我々のスコーピングレビューが強調しようとしている体系的問題です。')

    # Point 4
    add_para(doc, 'ポイント4 \u2013 GLP-1受容体作動薬:', bold=True, space_before=Pt(8))
    add_para(doc, 'レビュアーコメント: 現在許可されている。モニタリングリストに収載。禁止された場合、'
             '2型糖尿病管理などの正当な医学的理由があればTUEオプションが存在するだろう。',
             italic=True, indent=True)

    add_para(doc, '回答:', bold=True, space_before=Pt(4))
    add_para(doc, 'GLP-1受容体作動薬が現在許可されていること、禁止された場合にTUE経路が確立されるであろうことに'
             '同意します。我々の原稿はこの点を明示的に認めています。懸念は前向きなものです：禁止への軌道（2024年以降の'
             'モニタリングプログラム収載）を踏まえ、禁止後の事後的TUE枠組み構築ではなく、禁止前の積極的影響評価を'
             '提言しています。レビュアーの、禁止は体重階級制・審美系スポーツに限定される可能性が高いという指摘は、'
             'SCJ改訂版に反映する適切な示唆です。')

    # Point 5
    add_para(doc, 'ポイント5 \u2013 TUEガイドライン審査プロセス:', bold=True, space_before=Pt(8))
    add_para(doc, 'レビュアーコメント: 正式な審査プロセスが存在する。ほとんどのガイドラインは過去2年以内に更新済み。'
             '未完了は睡眠障害のレビューのみ。',
             italic=True, indent=True)

    add_para(doc, '回答:', bold=True, space_before=Pt(4))
    add_para(doc, 'この明確化に感謝し、WADAがTUE医師ガイドラインの更新に進展していることを認めます。臨床ガイドライン'
             '改訂とTUEガイドライン審査の間の正式な連携メカニズムに関する我々の提言は、既存プロセスの補完を意図して'
             'おり、置き換えを目指すものではありません。SCJ原稿では、最近の更新努力についてよりバランスの取れた言及を'
             '含めています。')

    # ---- Concluding remarks ----
    add_para(doc, '結語', bold=True, size=Pt(12), space_before=Pt(16))

    add_para(doc, '両査読者は主として規制側の視点からこのトピックにアプローチしていると思われます。'
             'これは貴重ですが、多面的な問題の一側面を代表するものです。我々の原稿は、日本で競技ボディビルダーとして'
             'アンチ・ドーピング制度を個人的に経験し、同様の課題に直面する患者を治療する医師─アスリートの二重の視点から'
             '執筆されています。この経験的視点は、規制枠組みだけからは明らかにならない実践的障壁の同定に寄与しています。')

    add_para(doc, '根本的な点を強調したいと思います：TUE経路の存在それ自体は、エビデンスに基づく医療への公平なアクセスを'
             '保証するものではありません。現行制度はアスリートに対し、標準治療を受けるために複雑な規制をナビゲートし、'
             '専門家の文書を確保し、例外を請願する負担を課しています。アンチ・ドーピング規制と自身の疾患の薬物療法の'
             '両方に精通したスポーツドクターにアクセスできるアスリートは限られています；大多数はWADA規則に不慣れな'
             '一般開業医や専門医に治療されています。我々は、標準的臨床治療が全てのアスリートにとっての基本的権利で'
             'あるべきであり、制限は物質が治療的回復を超える明確なパフォーマンス向上効果をもたらす場合にのみ適用される'
             'べきだと主張します。アスリートに課される治療ストレス──規制上の不確実性、行政的負担、これらの制度をナビゲート'
             'する際の心理的苦痛──は、それ自体が我々のレビューがより広い注目を集めることを目的とする重大かつ十分に'
             '認識されていない害です。')

    add_para(doc, 'Strength and Conditioning Journalへの投稿を進めます。完全なスコーピングレビュー形式'
             '（5,000〜8,000語）により、両査読者のコメントが示唆する包括的な取り扱いが可能です。SCJ原稿には、'
             'サルブタモール分析の拡充、より詳細なGLP-1議論、最近のTUEガイドライン更新の認知、そして──重要なことに──'
             '著者の実体験に基づくTUEアクセスの管轄間変動に関する新規セクションを含めています。')

    add_para(doc, '敬具', space_before=Pt(16))
    add_para(doc, '[著者名], MD, CSCS')

    output = os.path.join(OUTPUT_DIR, 'BJSM_Response_Letter_Japanese.docx')
    doc.save(output)
    print(f"Saved: {output}")


if __name__ == '__main__':
    create_english()
    create_japanese()
    print("Both response letters created.")
